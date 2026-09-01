"""P3a: MRR Word-document assembly (python-docx; no DB, no network)."""

import html
import io
import re

import pytest
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.services.reporting import (
    CONCLUSION,
    DOCX_MIMETYPE,
    REVIEW_HEADING,
    SUMMARY_INTRO,
    TITLE_SEPARATOR,
    UNDATED_LABEL,
    build_mrr_document,
    date_label,
    intro_sentence,
)


def test_build_mrr_document_saves():
    entries = [
        {"summaryDate": "01/02/2020", "summaryTitle": "Report A", "summaryText": "text A"},
        {"summaryDate": "03/04/2019", "summaryTitle": "Report B", "summaryText": "text B"},
    ]
    doc = build_mrr_document(
        entries,
        num_pages=42,
        patient_name="Synthetic Patient",
        patient_dob="-",
        qme_or_ame="QME",
        lawfirm="Example Law Firm",
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    assert buffer.tell() > 0
    assert DOCX_MIMETYPE.endswith("wordprocessingml.document")


def _intro(doc) -> str:
    """The 'I have received N pages...' sentence, wherever it sits in the document."""
    return next(p.text for p in doc.paragraphs if p.text.startswith("I have received"))


def test_intro_names_the_law_firm_when_there_is_one():
    doc = build_mrr_document(
        [], num_pages=8, patient_name="", patient_dob="", qme_or_ame="", lawfirm="Example Law Firm"
    )
    assert _intro(doc) == (
        "I have received 8 pages of medical records from Example Law Firm. I have reviewed all "
        "of the pages received and my opinion is based upon such received records."
    )


def test_intro_drops_the_clause_when_the_law_firm_is_blank():
    """A blank firm must not ship "medical records from ." into the delivered document.

    The field is optional free text on the review page and is routinely empty, so the old
    unconditional concatenation put a dangling preposition and an orphan full stop in front of every
    such export. Seen in a real one on 2026-08-17. An absent element is left out, which is the
    convention the title prompt already applies to a missing author.

    Asserted on both the presence of the correct sentence and the absence of the broken fragment,
    because a future edit could satisfy one without the other.
    """
    for blank in ("", "   ", None):
        doc = build_mrr_document(
            [], num_pages=8, patient_name="", patient_dob="", qme_or_ame="", lawfirm=blank
        )
        text = _intro(doc)
        assert text == (
            "I have received 8 pages of medical records. I have reviewed all of the pages "
            "received and my opinion is based upon such received records."
        )
        assert "from ." not in text
        assert "records from" not in text


def test_intro_has_no_double_space():
    """ "all of the pages  received" carried a double space in the shipped template."""
    doc = build_mrr_document(
        [], num_pages=1, patient_name="", patient_dob="", qme_or_ame="", lawfirm="Firm"
    )
    assert "  " not in _intro(doc)


def test_build_mrr_document_blank_qme_ame_does_not_crash():
    # A blank QME/AME field must not crash (an empty paragraph has no runs -> guarded with " ").
    doc = build_mrr_document(
        [], num_pages=1, patient_name="", patient_dob="", qme_or_ame="", lawfirm=""
    )
    buffer = io.BytesIO()
    doc.save(buffer)
    assert buffer.tell() > 0


def test_build_mrr_document_renders_two_column_table():
    # Entries render as a borderless 2-column table (date | title + text), sorted chronologically.
    entries = [
        {"summaryDate": "01/02/2020", "summaryTitle": "Report A", "summaryText": "text A"},
        {"summaryDate": "03/04/2019", "summaryTitle": "Report B", "summaryText": "text B"},
    ]
    doc = build_mrr_document(
        entries,
        num_pages=2,
        patient_name="Synthetic Patient",
        patient_dob="-",
        qme_or_ame="QME",
        lawfirm="Example Law Firm",
    )
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.columns) == 2
    assert len(table.rows) == 2
    # 03/04/2019 sorts before 01/02/2020; left cell = date, right cell = title + text.
    assert table.rows[0].cells[0].text == "03/04/2019"
    assert "Report B" in table.rows[0].cells[1].text
    assert "text B" in table.rows[0].cells[1].text


def test_summary_body_is_justified():
    # A finished report reads as justified prose; the date column stays left-aligned (default).
    entries = [
        {"summaryDate": "01/02/2020", "summaryTitle": "Report A", "summaryText": "text A"},
        {"summaryDate": "03/04/2019", "summaryTitle": "Report B", "summaryText": "text B"},
    ]
    doc = build_mrr_document(
        entries,
        num_pages=2,
        patient_name="Synthetic Patient",
        patient_dob="-",
        qme_or_ame="QME",
        lawfirm="Example Law Firm",
    )
    for row in doc.tables[0].rows:
        assert row.cells[1].paragraphs[0].alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        assert row.cells[0].paragraphs[0].alignment != WD_PARAGRAPH_ALIGNMENT.JUSTIFY


# --------------------------------------------------------------------------------------------------
# Both renderers, one sentence. `reporting` builds the Word document and `linked_pdf` builds the same
# letter as HTML for the combined PDF; each used to carry its own copy of every client-facing string.
# The copies had drifted before anyone noticed - the double space in "pages  received" was in the Word
# copy alone - so these pin the two renderers to each OTHER rather than each to its own expected
# literal. A test that only checks each renderer against a hardcoded string passes happily while the
# two disagree, which is exactly how the delivered .docx and .pdf came to differ.


def _pdf_letter(num_pages, lawfirm) -> str:
    """The linked-PDF letter HTML. Imported lazily so a missing pymupdf skips instead of erroring."""
    linked_pdf = pytest.importorskip("app.services.linked_pdf")
    return linked_pdf._summary_html([], num_pages, "Synthetic Patient", "-", "QME", lawfirm)


@pytest.mark.parametrize("lawfirm", ["Example Law Firm", "Smith & Jones, LLP", "", "   ", None])
def test_both_renderers_emit_the_same_intro_sentence(lawfirm):
    """The one sentence, identical in both artifacts, for every shape the firm field arrives in.

    The PDF side is compared against the HTML-ESCAPED sentence because that renderer escapes the
    assembled string on its way into the letter markup - which is why "Smith & Jones, LLP" is in the
    parameter list. Same sentence, escaped for its medium.
    """
    expected = intro_sentence(8, lawfirm)
    doc = build_mrr_document(
        [], num_pages=8, patient_name="", patient_dob="", qme_or_ame="", lawfirm=lawfirm
    )
    assert _intro(doc) == expected
    assert html.escape(expected) in _pdf_letter(8, lawfirm)


@pytest.mark.parametrize("lawfirm", ["", "   ", None])
def test_neither_renderer_ships_the_orphan_clause(lawfirm):
    """A blank firm must not put "medical records from ." in EITHER deliverable.

    The original fix only reached the Word document. The linked PDF interpolated the firm
    unconditionally, so a merge would have left the combined PDF - the artifact the defect was
    actually spotted in - still emitting it.
    """
    doc = build_mrr_document(
        [], num_pages=8, patient_name="", patient_dob="", qme_or_ame="", lawfirm=lawfirm
    )
    letter = _pdf_letter(8, lawfirm)
    for text in (_intro(doc), letter):
        assert "from ." not in text
        assert "records from" not in text


def test_both_renderers_share_the_other_two_sentences():
    """The intro was the one with the bug; these two are the remaining copies of the mechanism."""
    doc = build_mrr_document(
        [], num_pages=8, patient_name="", patient_dob="", qme_or_ame="", lawfirm="Firm"
    )
    paragraphs = [p.text for p in doc.paragraphs]
    letter = _pdf_letter(8, "Firm")
    for sentence in (SUMMARY_INTRO, CONCLUSION):
        assert sentence in paragraphs
        assert html.escape(sentence) in letter


# Undated entries, per the reviewers 2026-08-21: "if it is something important we will still
# summarize it, it can go at the end of the Review as Undated". Both renderers had it backwards - the
# sort key was datetime.min, so a document stating NO date sorted ahead of the earliest real
# encounter and the deliverable could OPEN on one - and the date cell showed the raw "-" the field
# spec writes, which reads as a value nobody filled in rather than a fact about the document.
def _entry(date, title="A REPORT", text="body text"):
    return {"summaryDate": date, "summaryTitle": title, "summaryText": text, "linkTitle": title}


@pytest.mark.parametrize("undated", ["-", "", "n/a", "   "])
def test_undated_entries_sort_last_in_the_word_document(undated):
    entries = [_entry(undated, "UNDATED"), _entry("01/02/2020", "EARLIEST"), _entry("03/04/2021")]
    doc = build_mrr_document(entries, 10, "A B", "01/01/1980", "QME", "Firm")
    # One borderless table, one row per entry, NO header row - see
    # test_build_mrr_document_renders_two_column_table. Indexing from row 1 silently skips the
    # first entry, which is how the first version of this test passed for the wrong reason.
    dates = [row.cells[0].text for row in doc.tables[0].rows]

    assert dates == ["01/02/2020", "03/04/2021", UNDATED_LABEL]


@pytest.mark.parametrize("undated", ["-", "", "   ", "n/a", "unknown"])
def test_the_date_cell_says_undated_rather_than_a_dash(undated):
    """One definition of undated, shared by the sort and the label. Written separately they
    disagreed: "n/a" sorted last but rendered as "n/a"."""
    assert date_label({"summaryDate": undated}) == UNDATED_LABEL


def test_a_missing_key_is_undated_rather_than_a_crash():
    assert date_label({}) == UNDATED_LABEL


def test_a_real_date_is_left_exactly_as_written():
    """Copy dates EXACTLY - the factuality rules say so, and a reviewer compares them to the page."""
    assert date_label({"summaryDate": "01/02/2020"}) == "01/02/2020"


def test_both_renderers_share_the_label_and_the_ordering():
    """They produce the same deliverable in two formats and a reviewer compares them side by side, so
    the wording and the ordering must not drift. Same function object, not merely equal behaviour."""
    from datetime import datetime

    from app.services import linked_pdf

    assert linked_pdf.date_label is date_label
    assert linked_pdf._sort_key(_entry("-")) == datetime.max
    assert linked_pdf._sort_key(_entry("01/02/2020")) < datetime.max


# The two client-facing sentences below are formatted by the LAST block of `build_mrr_document`,
# which set every property on the WRONG paragraph's run: `nine_title_format = fourth_title.runs[0]`.
# Two visible consequences in the delivered .docx, and ruff had flagged the cause as F841 (a local
# assigned and never used) before it was silenced with a noqa instead of fixed.
def _paragraph_named(doc, text):
    return next(p for p in doc.paragraphs if p.text == text)


def test_the_summary_intro_stays_bold_in_the_word_document():
    """It is set bold, then the conclusion block un-bolds it by reusing the same run.

    `linked_pdf` renders this sentence with `font-weight:bold`, so the delivered .docx and the
    delivered .pdf disagreed on the formatting of a sentence the client reads - the same drift the
    module docstring records for the sentence TEXT, one layer down in the formatting.
    """
    doc = build_mrr_document(
        [],
        num_pages=8,
        patient_name="A B",
        patient_dob="01/01/1980",
        qme_or_ame="QME",
        lawfirm="Firm",
    )
    runs = _paragraph_named(doc, SUMMARY_INTRO).runs
    assert runs, "the summary-intro paragraph has no runs"
    assert all(r.bold for r in runs), "the summary intro is not bold in the Word document"


def test_the_conclusion_is_formatted_like_the_rest_of_the_letter():
    """It got NO formatting, so the last sentence shipped in python-docx's default Calibri 11
    while every other paragraph is Times New Roman."""
    doc = build_mrr_document(
        [],
        num_pages=8,
        patient_name="A B",
        patient_dob="01/01/1980",
        qme_or_ame="QME",
        lawfirm="Firm",
    )
    runs = _paragraph_named(doc, CONCLUSION).runs
    assert runs, "the conclusion paragraph has no runs"
    for run in runs:
        assert run.font.name == "Times New Roman"
        assert run.font.size == Pt(12)
        assert not run.bold
        assert not run.underline


def test_paragraph_alignment_is_set_on_paragraphs_not_runs():
    """`alignment` is a paragraph property; assigning it to a run is silently a no-op.

    Two of these blocks set it on the run, so those paragraphs never had their alignment applied.
    Both happen to want LEFT, which python-docx also gives by default, so nothing was visible - but
    the next paragraph that wants CENTER would fail the same way and look like a docx quirk.
    """
    doc = build_mrr_document(
        [],
        num_pages=8,
        patient_name="A B",
        patient_dob="01/01/1980",
        qme_or_ame="QME",
        lawfirm="Firm",
    )
    for sentence in (SUMMARY_INTRO, CONCLUSION):
        paragraph = _paragraph_named(doc, sentence)
        assert paragraph.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT, (
            f"alignment never reached the paragraph for: {sentence!r}"
        )


def test_both_renderers_agree_that_the_summary_intro_is_bold():
    """The invariant that actually broke: not "is it bold in Word" but "do the two artifacts agree".

    Each renderer was individually self-consistent - the PDF hardcodes `font-weight:bold`, the Word
    side intended bold - and they still disagreed, because a later block in the Word assembly reused
    the wrong run and cleared it. Asserting the two sides together is what makes that visible.
    """
    doc = build_mrr_document(
        [],
        num_pages=8,
        patient_name="A B",
        patient_dob="01/01/1980",
        qme_or_ame="QME",
        lawfirm="Firm",
    )
    word_is_bold = all(r.bold for r in _paragraph_named(doc, SUMMARY_INTRO).runs)

    letter = _pdf_letter(8, "Firm")
    escaped = html.escape(SUMMARY_INTRO)
    paragraph = next(
        block for block in re.findall(r"<p[^>]*>.*?</p>", letter, re.S) if escaped in block
    )
    pdf_is_bold = "font-weight:bold" in paragraph.replace(" ", "")

    assert word_is_bold == pdf_is_bold, (
        f"the delivered .docx and .pdf disagree on whether {SUMMARY_INTRO!r} is bold "
        f"(word={word_is_bold}, pdf={pdf_is_bold})"
    )
    assert word_is_bold, "both renderers agree, but on NOT bold - the intended style is bold"


# The heading and the title separator were the two strings each renderer still held its own copy of,
# and the copies had already diverged. Both are checked against the eight human deliverables on disk:
# "MEDICAL RECORD REVIEW" in 8 of 8 files, and a PERIOD after the title in 329 of 329 date-anchored
# entries with not one colon. The Word renderer was wrong on both, the PDF renderer right on both.
def test_the_review_heading_is_the_form_the_human_deliverables_use():
    """8 of 8 human deliverables write it in caps. The Word renderer said "Medical Record Review"."""
    assert REVIEW_HEADING == "MEDICAL RECORD REVIEW"
    doc = build_mrr_document(
        [],
        num_pages=8,
        patient_name="A B",
        patient_dob="01/01/1980",
        qme_or_ame="QME",
        lawfirm="Firm",
    )
    paragraphs = [p.text for p in doc.paragraphs]
    assert REVIEW_HEADING in paragraphs
    assert "Medical Record Review" not in paragraphs, "the title-case heading is back"


def test_both_renderers_use_the_same_review_heading():
    doc = build_mrr_document(
        [],
        num_pages=8,
        patient_name="A B",
        patient_dob="01/01/1980",
        qme_or_ame="QME",
        lawfirm="Firm",
    )
    assert REVIEW_HEADING in [p.text for p in doc.paragraphs]
    assert html.escape(REVIEW_HEADING) in _pdf_letter(8, "Firm")


def test_the_title_is_separated_from_the_body_by_a_period_not_a_colon():
    """329 of 329 date-anchored human entries use a period. The Word renderer emitted ": "."""
    assert TITLE_SEPARATOR == ". "
    doc = build_mrr_document(
        [_entry("01/02/2020", "A REPORT", "body text")], 10, "A B", "01/01/1980", "QME", "Firm"
    )
    cell = doc.tables[0].rows[0].cells[1].text
    assert cell.startswith(f"A REPORT{TITLE_SEPARATOR}"), cell
    assert "A REPORT:" not in cell, "the colon separator is back"


def test_both_renderers_use_the_same_title_separator():
    entries = [_entry("01/02/2020", "A REPORT", "body text")]
    doc = build_mrr_document(entries, 10, "A B", "01/01/1980", "QME", "Firm")
    word_cell = doc.tables[0].rows[0].cells[1].text
    linked_pdf = pytest.importorskip("app.services.linked_pdf")
    letter = linked_pdf._summary_html(entries, 10, "A B", "01/01/1980", "QME", "Firm")
    assert word_cell.startswith(f"A REPORT{TITLE_SEPARATOR}")
    assert f"</a>{html.escape(TITLE_SEPARATOR)}" in letter


# One entry and one call, so the header tests below read as assertions about the header rather than
# as six copies of the same six-argument call.
_ENTRY = [{"summaryDate": "01/02/2026", "summaryTitle": "Report A", "summaryText": "text A"}]


def _build(entries):
    return build_mrr_document(
        entries,
        num_pages=42,
        patient_name="Doe, Jane",
        patient_dob="01/10/1961",
        qme_or_ame="QME",
        lawfirm="Example Law Firm",
    )


# ---------------------------------------------------------------------------------------------
# The running header. Three defects, all in the .docx - which is the PRIMARY deliverable - and all
# found by rendering both artifacts and diffing them, the same way the letter sentences were.


def _headers(doc):
    """(first-page header, later-page header) for the document's only section."""
    section = doc.sections[0]
    return section.first_page_header, section.header


def test_the_word_header_labels_the_date_of_birth():
    """WHEN the header names the patient, THE SYSTEM SHALL label the date of birth.

    The Word renderer wrote it as a bare value while the linked PDF wrote `DOB:`, so the two
    deliverables identified the patient differently on every page. `RE:` was labelled in both, which
    is what makes the unlabelled one an omission rather than a house style.
    """
    doc = _build(_ENTRY)
    first, later = _headers(doc)
    for header in (first, later):
        assert "DOB: 01/10/1961" in header.paragraphs[0].text


def test_the_word_header_carries_a_real_page_number_field():
    """IT SHALL number the pages, rather than writing the label and no number.

    The header ended with the literal string "Page " and nothing after it: python-docx writes text,
    not fields, so every page of the delivered .docx showed a dangling label while the linked PDF
    numbered its pages. A field is the only mechanism available - a Word header is ONE object
    repeated on every page, so there is no per-page text to write.
    """
    from docx.oxml.ns import qn

    doc = _build(_ENTRY)
    _first, later = _headers(doc)
    fields = later.paragraphs[0]._p.findall(qn("w:fldSimple"))
    assert [f.get(qn("w:instr")) for f in fields] == [" PAGE "]


def test_page_one_carries_the_patient_lines_and_no_number():
    """The first page SHALL show the identifying lines WITHOUT a page number, matching the linked
    PDF, which writes the number only from page 2 (`"" if i == 0`). A Word header applies to every
    page, so this needs a separate first-page header - which is also what stops "Page 1" appearing
    on a letter that opens with the patient's name."""
    from docx.oxml.ns import qn

    doc = _build(_ENTRY)
    assert doc.sections[0].different_first_page_header_footer is True
    first, _later = _headers(doc)
    assert "RE: Doe, Jane" in first.paragraphs[0].text
    assert "Page" not in first.paragraphs[0].text
    assert first.paragraphs[0]._p.findall(qn("w:fldSimple")) == []


def test_the_header_adds_no_blank_line_above_the_patient():
    """IT SHALL write into the header's existing paragraph rather than adding one.

    A new Word header already carries an empty paragraph, so `header.add_paragraph` left a blank
    line above the patient's name on every page of the deliverable. Asserted on the paragraph COUNT
    because that is the defect; the text assertions above pass either way.
    """
    doc = _build(_ENTRY)
    first, later = _headers(doc)
    assert len(first.paragraphs) == 1
    assert len(later.paragraphs) == 1
    assert first.paragraphs[0].text.startswith("RE:")
    assert later.paragraphs[0].text.startswith("RE:")


def test_both_renderers_take_the_header_lines_from_one_place():
    """WHEN either renderer names the patient, THE SYSTEM SHALL use the same two lines.

    The pin that matters: this module's docstring records that the letter's sentences drifted
    because each renderer held its own copy, and the header had drifted the same way for the same
    reason. Asserting each side separately would let them diverge again, so this asserts they AGREE.
    """
    from app.services.reporting import header_lines

    re_line, dob_line = header_lines("Doe, Jane", "01/10/1961")
    doc = _build(_ENTRY)
    _first, later = _headers(doc)
    assert re_line in later.paragraphs[0].text
    assert dob_line in later.paragraphs[0].text
    # and the linked PDF builds its running header from the same call
    import inspect

    from app.services import linked_pdf

    assert "header_lines(" in inspect.getsource(linked_pdf._draw_running_header)


def test_the_header_stays_times_new_roman_at_ten_point():
    """Unchanged by the rewrite: the header is 10pt Times New Roman, smaller than the 11pt body, and
    the linked PDF draws its own at fontsize=10. Pinned because the rewrite moved the run creation
    into a helper whose default size is the BODY's 11pt."""
    from docx.shared import Pt

    doc = _build(_ENTRY)
    _first, later = _headers(doc)
    run = later.paragraphs[0].runs[0]
    assert run.font.size == Pt(10)
    assert run.font.name == "Times New Roman"


def test_fill_header_clears_whatever_the_paragraph_already_held():
    """WHEN the header's first paragraph already carries runs, THE SYSTEM SHALL remove all of them.

    `_fill_header` reuses the existing paragraph rather than adding one, so anything already in it
    has to go or it would sit above the patient line on every page. Nothing else reaches this loop:
    a fresh header's first paragraph has no runs, which is why the removal was uncovered until now.

    It also guards the `list()` that used to wrap the loop. Deleting the elements being iterated is
    only safe because python-docx rebuilds `paragraph.runs` from the XML on each access; if that
    ever stops being true, this test fails rather than the deliverable silently keeping stale runs.
    """
    import docx

    from app.services.reporting import _fill_header

    document = docx.Document()
    header = document.sections[0].header
    header.paragraphs[0].add_run("stale name")
    header.paragraphs[0].add_run(" and a stale dob")
    assert len(header.paragraphs[0].runs) == 2

    _fill_header(header, "RE: Synthetic Patient", "DOB: 01/01/1990", numbered=False)

    text = header.paragraphs[0].text
    assert "stale" not in text
    assert "Synthetic Patient" in text
    assert "01/01/1990" in text
