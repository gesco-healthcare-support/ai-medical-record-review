"""MRR Word-document assembly (ported from the Flask export blueprint; python-docx, Flask-free).

`build_mrr_document` is shared by the export + bundle-summarize routes: it sorts summary entries
chronologically and assembles the letterhead + intro + per-record body into a python-docx
Document. The classic CSV/on-disk export routes are dropped.
"""

import re
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# The three client-facing sentences of the letter live HERE ONLY, and both renderers import them:
# this module builds the Word document, `linked_pdf` builds the same letter as HTML for the combined
# PDF. They used to hold their own copy of each string, and the copies had already drifted - the
# double space in "pages  received" was in the Word copy alone, so the delivered .docx and the
# delivered .pdf disagreed on a sentence the client reads. Two copies of one sentence silently
# diverging is the mechanism, not the typo, so there is one home for the text and a test that pins
# both renderers to it.
SUMMARY_INTRO = "The following is a summary of those records:"
CONCLUSION = "This concludes the review of submitted records."

# Same reason as the three sentences above: both renderers held their own copy and the copies had
# ALREADY diverged. Checked against the eight human deliverables on disk, which are the standard these
# artifacts are supposed to match:
#
#   heading    "MEDICAL RECORD REVIEW" in 8 of 8 files. The Word renderer said "Medical Record
#              Review"; the PDF renderer already had it in caps.
#   separator  329 date-anchored entries across those 8 files, every one a PERIOD after the title
#              and not one colon. The Word renderer emitted ": "; the PDF renderer already used ". ".
#
# So on both counts the .docx disagreed with the .pdf AND with the human standard, and the .docx is
# the primary deliverable. Centralised here so the next change moves both renderers at once.
REVIEW_HEADING = "MEDICAL RECORD REVIEW"
TITLE_SEPARATOR = ". "
_REPORT_FONT = "Times New Roman"


def header_lines(patient_name, patient_dob) -> tuple[str, str]:
    """The two identifying lines both renderers put at the top of every page.

    Here for the same reason as the sentences above, and it had already gone the same way: the Word
    header wrote the date of birth as a bare value while the linked PDF labelled it ``DOB:``, so the
    two deliverables named the patient differently on every page. `RE:` was labelled in both, which
    is what makes the unlabelled one read as an omission rather than a house style.
    """
    return f"RE: {patient_name}", f"DOB: {patient_dob}"


def intro_sentence(num_pages, lawfirm) -> str:
    """The letter's opening sentence, shared by the Word and linked-PDF renderers.

    The law firm is OPTIONAL - free text on the review page that a reviewer often has no value for.
    Concatenating it unconditionally shipped "medical records from ." into the delivered document,
    dangling preposition and orphan full stop, every time the field was blank; seen in a real export
    on 2026-08-17. So the clause is DROPPED rather than rendered empty - an absent element is left
    out, the convention the title prompt already applies to a missing author. ""/whitespace/None all
    count as absent, which is how the value actually arrives from the form.

    Returns plain text. The PDF renderer escapes the ASSEMBLED sentence for HTML rather than
    escaping `lawfirm` before it gets here, so the escape still covers any field added later.
    """
    firm = (lawfirm or "").strip()
    received_from = f" from {firm}" if firm else ""
    return (
        f"I have received {num_pages} pages of medical records{received_from}. "
        "I have reviewed all of the pages received and my opinion is based upon such "
        "received records."
    )


# Inline emphasis the summarizer emits: **bold**, *italic*, _italic_. Rendered as real runs so no
# raw markers leak into the Word document (mirrors the web MarkdownText renderer).
_INLINE_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_", re.DOTALL)


def _run(paragraph, s, *, bold=False, italic=False, size=None):
    run = paragraph.add_run(s)
    run.bold = bold
    run.italic = italic
    run.font.name = _REPORT_FONT
    run.font.size = size or Pt(11)
    return run


def _add_inline_runs(paragraph, text, *, bold=False, italic=False):
    """Append runs to ``paragraph``, turning **bold** / *italic* / _italic_ markers into real
    formatting; ``bold``/``italic`` set the baseline for the plain segments."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            _run(paragraph, text[pos : m.start()], bold=bold, italic=italic)
        if m.group(1) is not None:
            _run(paragraph, m.group(1), bold=True, italic=italic)
        else:
            _run(paragraph, m.group(2) or m.group(3), bold=bold, italic=True)
        pos = m.end()
    if pos < len(text):
        _run(paragraph, text[pos:], bold=bold, italic=italic)


def _page_number_field(paragraph) -> None:
    """Append a real Word PAGE field, so the header numbers the pages.

    The header used to end with the literal string ``"Page "`` and nothing after it: python-docx
    writes text, not fields, so every page of the delivered .docx carried a dangling label with no
    number while the linked PDF numbered its pages properly. A field rather than a computed integer
    because a Word header is one object repeated on every page - there is no per-page text to write,
    which is presumably how the bare label got there.

    ``w:fldSimple`` with a cached ``1`` inside: Word and LibreOffice both recalculate the field on
    open, and the cached run is what a reader that does not evaluate fields shows instead of nothing.
    """
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " PAGE ")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def _fill_header(header, re_line, dob_line, *, numbered: bool) -> None:
    """Write the identifying lines into ``header``, optionally followed by ``Page <n>``.

    Writes into the header's EXISTING first paragraph rather than adding one. A new Word header
    already carries an empty paragraph, so `add_paragraph` left a blank line above the patient's
    name on every page of the deliverable.
    """
    paragraph = header.paragraphs[0]
    # Removing the run ELEMENTS, not `paragraph.text = ""` - that assignment leaves one empty run
    # behind, so the first run of the header carried no font at all and `runs[0]` was not the text.
    # Safe to iterate while deleting: `runs` builds a fresh list from the XML on every access, so
    # this is already a snapshot and the copy that used to wrap it added nothing.
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    _run(paragraph, f"{re_line}\n{dob_line}" + ("\nPage " if numbered else ""), size=Pt(10))
    if numbered:
        _page_number_field(paragraph)


UNDATED_LABEL = "Undated"


def parsed_date(entry):
    """The entry's date as a datetime, or None when it states none.

    ONE definition of "undated", shared by the sort key and the label. Written separately they
    disagreed: an entry dated "n/a" sorted last (unparseable) but rendered the raw "n/a", because the
    label only special-cased "-" and "". So an entry could sort as undated while displaying junk. Not
    reachable from today's field spec, which only ever writes "-", but two definitions of one concept
    drift eventually.
    """
    try:
        return datetime.strptime((entry.get("summaryDate") or "").strip(), "%m/%d/%Y")
    except ValueError:
        return None


def date_label(entry) -> str:
    """The date cell text: the date exactly as written, or "Undated".

    The field spec writes "-" when a document carries no date, which in a finished deliverable reads
    as a value nobody filled in rather than a fact about the document. The reviewers call these
    Undated and expect them at the end of the review.

    A parsed date is returned VERBATIM, never reformatted: the factuality rules say copy dates
    exactly, and a reviewer compares them against the page.
    """
    if parsed_date(entry) is None:
        return UNDATED_LABEL
    return (entry.get("summaryDate") or "").strip()


def build_mrr_document(entries, num_pages, patient_name, patient_dob, qme_or_ame, lawfirm):
    """Assemble the MRR Word document from summary ``entries`` (sorted chronologically)."""

    # Undated entries sort LAST, per the reviewers 2026-08-21: "if it is something important we will
    # still summarize it, it can go at the end of the Review as Undated". This was datetime.min -
    # undated FIRST - so a document stating no date sorted ahead of the earliest real encounter and
    # the deliverable could OPEN on one.
    entries = sorted(entries, key=lambda e: parsed_date(e) or datetime.max)

    doc = Document()

    # HEADER
    section = doc.sections[0]
    # A separate first-page header, so page 1 carries the patient lines WITHOUT a page number and
    # every later page carries all three. That is what `linked_pdf._draw_running_header` already
    # does (`"" if i == 0 else f"\nPage {i + 1}"`), and matching it is the point: the two are the
    # same deliverable in two formats and a reviewer reads them side by side.
    section.different_first_page_header_footer = True
    re_line, dob_line = header_lines(patient_name, patient_dob)
    _fill_header(section.first_page_header, re_line, dob_line, numbered=False)
    _fill_header(section.header, re_line, dob_line, numbered=True)

    doc.add_paragraph("")

    # TITLE. An empty paragraph has no runs, so runs[0] below would crash on a blank
    # QME/AME field (the form allows leaving it empty).
    title = doc.add_paragraph(qme_or_ame or " ")
    title_format = title.runs[0]
    title_format.bold = True
    title_format.underline = True
    title_format.font.size = Pt(12)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_format.font.name = _REPORT_FONT

    doc.add_paragraph("")

    second_title = doc.add_paragraph(REVIEW_HEADING)
    second_title_format = second_title.runs[0]
    second_title_format.bold = True
    second_title_format.underline = True
    second_title_format.font.size = Pt(12)
    second_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    second_title_format.font.name = _REPORT_FONT

    intro_text = intro_sentence(num_pages, lawfirm)
    second_intro_text = SUMMARY_INTRO
    this_concludes_text = CONCLUSION

    third_title = doc.add_paragraph(intro_text)
    third_title_format = third_title.runs[0]
    third_title_format.bold = False
    third_title_format.underline = False
    third_title_format.font.size = Pt(12)
    third_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    third_title_format.font.name = _REPORT_FONT

    fourth_title = doc.add_paragraph(second_intro_text)
    fourth_title_format = fourth_title.runs[0]
    fourth_title_format.bold = True
    fourth_title_format.underline = False
    fourth_title_format.font.size = Pt(12)
    fourth_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    fourth_title_format.font.name = _REPORT_FONT

    # Two-column borderless table: date | title + body. The default "Table Normal" style has no
    # cell borders, matching the canonical MRR summary layout (date sits in its own left column,
    # the summary flows in the right column) instead of the old inline date-tab-title paragraph.
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for entry in entries:
        cells = table.add_row().cells
        cells[0].width = Inches(0.9)
        cells[1].width = Inches(5.6)
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _run(cells[0].paragraphs[0], date_label(entry))
        body = cells[1].paragraphs[0]
        # Justified: the report is read as a finished document, and a ragged right edge on every
        # record is what made the export look like a draft.
        body.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        _add_inline_runs(body, entry["summaryTitle"], bold=True)
        _run(body, TITLE_SEPARATOR)
        _add_inline_runs(body, entry["summaryText"])

    # `nine_title_format` read `fourth_title.runs[0]` - the SUMMARY_INTRO paragraph's run, not this
    # one. Two visible defects in the delivered .docx from one wrong name: the conclusion got no
    # formatting at all and shipped in python-docx's default Calibri 11 while every other paragraph
    # is Times New Roman 12, and `bold = False` here UNDID the `bold = True` set on SUMMARY_INTRO
    # thirty lines above. `linked_pdf` renders that sentence with `font-weight:bold`, so the .docx and
    # the .pdf disagreed on the formatting of a sentence the client reads - the same drift this
    # module's docstring records for the sentence TEXT, one layer down.
    #
    # ruff had already found it: `nine_title` was assigned and never used (F841), which is exactly
    # the bug, and the warning was silenced with a noqa rather than fixed.
    #
    # `alignment` is a PARAGRAPH property, so it is set on the paragraph here and on `fourth_title`
    # above. Assigning it to a run is silently a no-op; both wanted LEFT, which is also the default,
    # so nothing was visible - but the next paragraph wanting CENTER would fail the same way.
    nine_title = doc.add_paragraph(this_concludes_text)
    nine_title_format = nine_title.runs[0]
    nine_title_format.bold = False
    nine_title_format.underline = False
    nine_title_format.font.size = Pt(12)
    nine_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    nine_title_format.font.name = _REPORT_FONT

    return doc
