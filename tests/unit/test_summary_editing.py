"""Per-row summarize marking, summary editing/exclusion, and the boot migration.

All engines stubbed; synthetic data only.
"""

import io
import sqlite3
import time


def _upload(client, pdf_bytes):
    response = client.post(
        "/api/documents",
        data={"pdf": (io.BytesIO(pdf_bytes(10)), "synthetic.pdf")},
        content_type="multipart/form-data",
    )
    assert response.status_code == 201
    return response.get_json()["id"]


def _wait_done(client, document_id, timeout=10.0):
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        payload = client.get(f"/api/documents/{document_id}/status").get_json()
        if payload["status"] == "done":
            return True
        if payload["status"] == "error":
            raise AssertionError(f"job failed: {payload['job']}")
        time.sleep(0.02)
    return False


_ROWS = [
    {
        "start": 1,
        "end": 4,
        "category": "1",
        "title": "DOC A",
        "date": "01/15/2024",
        "injury_date": "-",
        "flag": "-",
        "include": True,
    },
    {
        "start": 5,
        "end": 7,
        "category": "3",
        "title": "DOC B",
        "date": "02/20/2024",
        "injury_date": "-",
        "flag": "-",
        "include": False,
    },
    {
        "start": 8,
        "end": 10,
        "category": "5",
        "title": "DOC C",
        "date": "03/05/2024",
        "injury_date": "-",
        "flag": "x",
        "include": True,
    },
]


def _fake_summarize(pdf_path, row, model=None):
    return {
        "summaryTitle": f"SUMMARY {row['start']}",
        "summaryDate": row["date"],
        "summaryText": "synthetic text",
        "manualCheck": row["flag"] == "x",
        "sourceText": f"extracted page text {row['start']}-{row['end']}",
    }


def test_only_included_rows_are_summarized(client, pdf_bytes, monkeypatch):
    monkeypatch.setattr("mrr_ai.services.summarize_engine.summarize_row", _fake_summarize)
    document_id = _upload(client, pdf_bytes)

    assert (
        client.post(
            f"/api/documents/{document_id}/summarize/start", json={"rows": _ROWS}
        ).status_code
        == 200
    )
    assert _wait_done(client, document_id)

    summaries = client.get(f"/api/documents/{document_id}/summaries").get_json()
    assert [item["summaryTitle"] for item in summaries] == ["SUMMARY 1", "SUMMARY 8"]

    detail = client.get(f"/api/documents/{document_id}").get_json()
    assert [row["include"] for row in detail["rows"]] == [True, False, True]


def test_summarize_refused_when_nothing_included(client, pdf_bytes):
    document_id = _upload(client, pdf_bytes)
    rows = [dict(row, include=False) for row in _ROWS]
    response = client.post(f"/api/documents/{document_id}/summarize/start", json={"rows": rows})
    assert response.status_code == 400
    assert "marked for summarization" in response.get_json()["error"]


def test_edit_and_exclude_summary(client, pdf_bytes, monkeypatch):
    monkeypatch.setattr("mrr_ai.services.summarize_engine.summarize_row", _fake_summarize)
    document_id = _upload(client, pdf_bytes)
    client.post(f"/api/documents/{document_id}/summarize/start", json={"rows": _ROWS})
    assert _wait_done(client, document_id)

    edited = client.put(
        f"/api/documents/{document_id}/summaries/0",
        json={"summaryText": "corrected by reviewer", "summaryTitle": "FIXED TITLE"},
    )
    assert edited.status_code == 200
    body = edited.get_json()
    assert body["summaryText"] == "corrected by reviewer"
    assert body["edited"] is True

    # The raw model output is preserved underneath (training signal).
    from mrr_ai.models import Summary

    with client.application.app_context():
        raw = Summary.query.filter_by(document_id=document_id, idx=0).one()
        assert raw.text == "synthetic text"
        assert raw.edited_text == "corrected by reviewer"
        # The model INPUT is captured too: (source_text, text, edited_text) is the
        # complete fine-tuning triple for the summarizer.
        assert raw.source_text == "extracted page text 1-4"

    assert (
        client.put(f"/api/documents/{document_id}/summaries/1", json={"excluded": True}).status_code
        == 200
    )
    listing = client.get(f"/api/documents/{document_id}/summaries").get_json()
    assert [item["excluded"] for item in listing] == [False, True]

    unknown = client.put(f"/api/documents/{document_id}/summaries/99", json={"excluded": True})
    assert unknown.status_code == 404


def test_export_skips_excluded_and_uses_edits(client, pdf_bytes, monkeypatch):
    monkeypatch.setattr("mrr_ai.services.summarize_engine.summarize_row", _fake_summarize)
    document_id = _upload(client, pdf_bytes)
    client.post(f"/api/documents/{document_id}/summarize/start", json={"rows": _ROWS})
    assert _wait_done(client, document_id)

    client.put(f"/api/documents/{document_id}/summaries/0", json={"excluded": True})
    exported = client.post(f"/api/documents/{document_id}/export", json={"QMEorAME": "QME"})
    assert exported.status_code == 200
    assert exported.data[:2] == b"PK"

    client.put(f"/api/documents/{document_id}/summaries/1", json={"excluded": True})
    all_excluded = client.post(f"/api/documents/{document_id}/export", json={})
    assert all_excluded.status_code == 409


def test_export_recomposes_tags_around_edits(client, pdf_bytes, monkeypatch):
    """The web UI edits CLEAN titles/text ([ManualCheck] and **DOI** stripped for
    display); the docx must still carry the legacy tag format for every summary."""

    def decorated(pdf_path, row, model=None):
        manual = "[ManualCheck] " if row["flag"] == "x" else ""
        return {
            "summaryTitle": f"{manual}SUMMARY {row['start']} (Pages {row['start']}-{row['end']})",
            "summaryDate": row["date"],
            "summaryText": "**DOI**:06/01/2023, original body",
            "manualCheck": row["flag"] == "x",
        }

    monkeypatch.setattr("mrr_ai.services.summarize_engine.summarize_row", decorated)
    document_id = _upload(client, pdf_bytes)
    client.post(f"/api/documents/{document_id}/summarize/start", json={"rows": _ROWS})
    assert _wait_done(client, document_id)

    # idx 1 is the flagged row (_ROWS[2], flag "x"); edit it the way the UI does: clean.
    assert (
        client.put(
            f"/api/documents/{document_id}/summaries/1",
            json={"summaryTitle": "EDITED TITLE (Pages 8-10)", "summaryText": "edited body"},
        ).status_code
        == 200
    )

    exported = client.post(f"/api/documents/{document_id}/export", json={"QMEorAME": "QME"})
    assert exported.status_code == 200

    from docx import Document as DocxDocument

    text = "\n".join(p.text for p in DocxDocument(io.BytesIO(exported.data)).paragraphs)
    assert "[ManualCheck] EDITED TITLE" in text  # flag recomposed around the clean edit
    assert "**DOI**:06/01/2023," in text  # DOI prefix recovered from the raw text
    assert "edited body" in text
    assert "[ManualCheck] [ManualCheck]" not in text  # unedited titles are not doubled


def test_export_recomposes_pages_and_diagnostic_tags(client, pdf_bytes, monkeypatch):
    """The web UI also strips the " (Pages X-Y)" suffix and the "[Diagnostic Study]"
    tag from titles; the docx must recompose BOTH from the summary's row snapshot,
    normalized to the canonical row range - regardless of what the reviewer typed."""

    rows = [
        dict(_ROWS[0]),  # category 1
        dict(_ROWS[2], category="3", flag="-"),  # category 3 -> [Diagnostic Study]
    ]

    def engine_style(pdf_path, row, model=None):
        diag = " [Diagnostic Study]" if str(row["category"]) == "3" else ""
        return {
            "summaryTitle": f"RAW TITLE{diag} (Pages {row['start']}-{row['end']})",
            "summaryDate": row["date"],
            "summaryText": "body text",
            "manualCheck": False,
        }

    monkeypatch.setattr("mrr_ai.services.summarize_engine.summarize_row", engine_style)
    document_id = _upload(client, pdf_bytes)
    client.post(f"/api/documents/{document_id}/summarize/start", json={"rows": rows})
    assert _wait_done(client, document_id)

    # The reviewer edits both titles CLEAN (the way the redesigned UI presents them);
    # the second edit even carries a stale typed suffix that must be normalized.
    client.put(f"/api/documents/{document_id}/summaries/0", json={"summaryTitle": "CLEAN EDIT"})
    client.put(
        f"/api/documents/{document_id}/summaries/1",
        json={"summaryTitle": "DIAG EDIT (Pages 99-99)"},
    )

    exported = client.post(f"/api/documents/{document_id}/export", json={"QMEorAME": "QME"})
    assert exported.status_code == 200

    from docx import Document as DocxDocument

    text = "\n".join(p.text for p in DocxDocument(io.BytesIO(exported.data)).paragraphs)
    assert "CLEAN EDIT (Pages 1-4)" in text
    assert "DIAG EDIT [Diagnostic Study] (Pages 8-10)" in text  # normalized, tag restored
    assert "(Pages 99-99)" not in text  # typed suffix replaced by the row snapshot
    assert "[Diagnostic Study] [Diagnostic Study]" not in text


def test_boot_migration_adds_columns_to_old_database(tmp_path):
    """A database created before the include/edited_*/excluded columns must upgrade in
    place on boot - the seeded demo data cannot be thrown away for a schema change."""
    db_path = tmp_path / "old.db"
    connection = sqlite3.connect(db_path)
    connection.execute(
        "CREATE TABLE review_rows (id INTEGER PRIMARY KEY, document_id VARCHAR(36),"
        ' idx INTEGER, start INTEGER, "end" INTEGER, category VARCHAR(8),'
        " title VARCHAR(512), date VARCHAR(16), injury_date VARCHAR(16),"
        " flag VARCHAR(4), suggest_merge BOOLEAN)"
    )
    connection.execute(
        "CREATE TABLE summaries (id INTEGER PRIMARY KEY, document_id VARCHAR(36),"
        " job_id INTEGER, idx INTEGER, title VARCHAR(512), date VARCHAR(16),"
        " text TEXT, manual_check BOOLEAN, row_start INTEGER, row_end INTEGER,"
        " row_category VARCHAR(8))"
    )
    connection.commit()
    connection.close()

    from mrr_ai import create_app

    create_app(
        {
            "TESTING": True,
            "SQLALCHEMY_DATABASE_URI": "sqlite:///" + str(db_path),
            "WTF_CSRF_ENABLED": False,
            "SECURITY_PASSWORD_HASH": "plaintext",
        }
    )

    connection = sqlite3.connect(db_path)
    review_columns = {row[1] for row in connection.execute("PRAGMA table_info(review_rows)")}
    summary_columns = {row[1] for row in connection.execute("PRAGMA table_info(summaries)")}
    connection.close()
    assert "include" in review_columns
    assert {
        "edited_title",
        "edited_date",
        "edited_text",
        "excluded",
        "source_text",
    } <= summary_columns
