"""Export routes: results CSV and the assembled MRR Word documents."""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from flask import Blueprint, jsonify, request, send_file

from mrr_ai import state
from mrr_ai.services.files import safe_name

bp = Blueprint("export", __name__)

_DOCX_MIMETYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_mrr_document(patient_name, patient_dob, qme_or_ame, lawfirm):
    """Assemble the MRR Word document from ``state.all_data`` (sorted chronologically).

    Shared by both Word-export routes; they differ only in where the document is saved
    and how it is returned, so the document-building logic lives here once.
    """

    def safe_date_parse(entry):
        try:
            return datetime.strptime(entry.get("summaryDate", ""), "%m/%d/%Y")
        except ValueError:
            return datetime.min  # Assign a very early date as fallback

    state.all_data = sorted(state.all_data, key=safe_date_parse)

    doc = Document()

    # HEADER
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.add_paragraph(
        "RE: " + patient_name + "\n" + patient_dob + "\n" + "Page "
    )
    for run in header_paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

    doc.add_paragraph("")

    # TITLE
    title = doc.add_paragraph(qme_or_ame)
    title_format = title.runs[0]
    title_format.bold = True
    title_format.underline = True
    title_format.font.size = Pt(12)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_format.font.name = "Times New Roman"

    doc.add_paragraph("")

    # Medical Record Review
    second_title = doc.add_paragraph("Medical Record Review")
    second_title_format = second_title.runs[0]
    second_title_format.bold = True
    second_title_format.underline = True
    second_title_format.font.size = Pt(12)
    second_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    second_title_format.font.name = "Times New Roman"

    intro_text = (
        "I have received "
        + str(state.num_pages)
        + " pages of medical records from "
        + lawfirm
        + ". I have reviewed all of the pages  received and my opinion is based upon such received records."
    )
    second_intro_text = "The following is a summary of those records:"
    this_concludes_text = "This concludes the review of submitted records."

    # First Line
    third_title = doc.add_paragraph(intro_text)
    third_title_format = third_title.runs[0]
    third_title_format.bold = False
    third_title_format.underline = False
    third_title_format.font.size = Pt(12)
    third_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    third_title_format.font.name = "Times New Roman"

    # Second Line
    fourth_title = doc.add_paragraph(second_intro_text)
    fourth_title_format = fourth_title.runs[0]
    fourth_title_format.bold = True
    fourth_title_format.underline = False
    fourth_title_format.font.size = Pt(12)
    fourth_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    fourth_title_format.font.name = "Times New Roman"

    big_text = ""
    main_paragraph = None
    for entry in state.all_data:
        big_text += (
            f"_{entry['summaryDate']}_\t****{entry['summaryTitle']}****: {entry['summaryText']}"
        )
        main_paragraph = doc.add_paragraph(big_text)
        big_text = ""

    if main_paragraph is not None:
        for run in main_paragraph.runs:  # Ensure the content uses Times New Roman
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)

    nine_title = doc.add_paragraph(this_concludes_text)  # noqa: F841
    nine_title_format = fourth_title.runs[0]
    nine_title_format.bold = False
    nine_title_format.underline = False
    nine_title_format.font.size = Pt(12)
    nine_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    nine_title_format.font.name = "Times New Roman"

    return doc


@bp.route("/exportresultstoCSV", methods=["POST"])
def exportresultstoCSV():
    try:
        # Get the JSON data from the request
        data = request.get_json()

        # Extract the text content and strip unwanted quotes
        csv_content = data.get("TXTText", "").strip('"')
        print(csv_content)

        # Prepare the file path
        file_path = os.path.join(os.path.expanduser("~"), "MRRs", state.main_filename + ".csv")
        print(state.main_filename)

        # Write the content to the CSV file without using csv.writer
        with open(file_path, "w", encoding="utf-8") as csv_file:
            csv_file.write(csv_content)

        # Return a success response
        return jsonify({"message": "Content saved to CSV file successfully!"}), 200
    except Exception as e:
        # Handle any errors
        return jsonify({"error": str(e)}), 500


@bp.route("/exportresultstoword", methods=["POST"])
def exportresultstoword():
    patientName = request.json.get("patientName")
    patientdob = request.json.get("patientdob")
    QMEorAME = request.json.get("QMEorAME")
    lawfirm = request.json.get("lawfirm")

    # Check if main_filename is None and set a default value if needed
    if state.main_filename is None:
        state.main_filename = "default_filename"

    doc = _build_mrr_document(patientName, patientdob, QMEorAME, lawfirm)

    out_dir = os.path.join(os.path.expanduser("~"), "MRRs")
    os.makedirs(out_dir, exist_ok=True)  # fresh machines lack ~/MRRs; export must not 500
    file_path_int = os.path.join(out_dir, state.main_filename + "_int.docx")
    file_path_rep = os.path.join(out_dir, state.main_filename + "_rep.docx")

    doc.save(file_path_int)
    response_int = send_file(file_path_int, as_attachment=True, mimetype=_DOCX_MIMETYPE)

    doc.save(file_path_rep)
    response_rep = send_file(file_path_rep, as_attachment=True, mimetype=_DOCX_MIMETYPE)

    response_int.headers["Content-Disposition"] = "attachment; filename=summaries.docx"
    response_rep.headers["Content-Disposition"] = "attachment; filename=summaries.docx"

    return response_int


@bp.route("/exportResultsToWordFileIndivRecords", methods=["POST"])
def exportResultsToWordFileIndivRecords():
    patientName = request.json.get("patientName")
    patientdob = request.json.get("patientdob")
    QMEorAME = request.json.get("QMEorAME")
    lawfirm = request.json.get("lawfirm")

    doc = _build_mrr_document(patientName, patientdob, QMEorAME, lawfirm)

    # Sanitize the patient name before using it in the output filename (prevents traversal).
    file_path = os.path.join(
        os.path.expanduser("~"), "MRRs", safe_name(patientName) + " - MRR.docx"
    )
    doc.save(file_path)

    response = send_file(file_path, as_attachment=True, mimetype=_DOCX_MIMETYPE)
    response.headers["Content-Disposition"] = "attachment; filename=summaries.docx"

    return response
