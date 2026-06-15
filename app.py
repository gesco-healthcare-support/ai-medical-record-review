import csv
import json
import os
import re
import time
from datetime import datetime
from difflib import SequenceMatcher

import httplib2
import pytesseract
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from dotenv import load_dotenv
from flask import Flask, jsonify, render_template, request, send_file
from google import genai
from google.genai import types
from openai import OpenAI
from pdf2image import convert_from_path
from pypdf import PdfReader, PdfWriter

from groups import groups
from prompts import prompts

# Set a higher timeout value
http = httplib2.Http(timeout=600)  # Set timeout to 5 minutes

load_dotenv()

# Fail fast if required secrets are not configured (see .env.example).
_required_env = ("GEMINI_API_KEY", "OPENAI_API_KEY")
_missing_env = [name for name in _required_env if not os.environ.get(name)]
if _missing_env:
    raise RuntimeError(
        "Missing required environment variables: "
        + ", ".join(_missing_env)
        + ". Copy .env.example to .env and fill in the values."
    )

genai_client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])

app = Flask(__name__)

app.config["MAX_CONTENT_LENGTH"] = 1024 * 1024 * 1024
script_dir = os.path.dirname(os.path.abspath(__file__))
uploads_dir = os.path.join(script_dir, "uploads")
app.config["UPLOAD_FOLDER"] = uploads_dir
ALLOWED_EXTENSIONS = {"pdf"}

pdf_filepath = None
txt_filepath = None
pdf_savepath = "/home/usera/mrr-line/uploads/"
main_filename = "summary"
main_txt_filename = "txt_pages"
patientNameGlobal = "Patient Full Name"
pages_not_counting = 0
num_pages = 0
all_data = []
manual_intervention = ""
indiv_mrr_folder_path = ""


# OPENAI_API_KEY is read from the environment by the OpenAI client (see .env.example).
client = OpenAI()


###########################
# Methods
###########################


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


# Function to clean and normalize strings for comparison
def normalize(text):
    return re.sub(r"[^a-zA-Z0-9\s]", "", text).strip().lower()


# Function to calculate similarity
def similarity(a, b):
    return SequenceMatcher(None, a, b).ratio()


def categorize_documents(title, categories, threshold=0.65):
    if not isinstance(title, str):
        print(f"Warning: Invalid title encountered: {title}")
        title = "Unknown"  # Fallback value for invalid or missing titles
    normalized_title = normalize(title)
    best_match = None
    best_group = None
    highest_similarity = 0

    # Check each category
    for group, docs in categories.items():
        for doc in docs:
            normalized_doc = normalize(doc)
            sim = similarity(normalized_title, normalized_doc)
            if sim > highest_similarity and sim >= threshold:
                highest_similarity = sim
                best_match = doc  # noqa: F841
                best_group = group

    # If no group is found with sufficient similarity, assign to "Group 100"
    if not best_group:
        best_group = "100"

    return best_group


def get_pdf_size(filepath):
    """Returns the size of the PDF file in megabytes."""
    file_size = os.path.getsize(filepath)  # Get size in bytes
    size_in_mb = file_size / (1024 * 1024)  # Convert bytes to megabytes
    return size_in_mb


def get_pdf_page_count(pdf_file):
    try:
        # Open the PDF file
        reader = PdfReader(pdf_file)
        # Get the total number of pages
        total_pages = len(reader.pages)
        return total_pages
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None


# def segment_pdf_2(pdf_path, max_pages=250):
#     """
#     Segments a PDF file into multiple PDF files, each containing a maximum of `max_pages` pages.

#     Parameters:
#         pdf_path (str): The path to the PDF file.
#         max_pages (int): Maximum number of pages in each segmented PDF file. Default is 250 pages.

#     Returns:
#         None
#     """
#     if not os.path.exists(pdf_path):
#         print(f"Error: File '{pdf_path}' does not exist.")
#         return

#     # Extract the main filename without extension
#     pdf_basename = os.path.splitext(os.path.basename(pdf_path))[0]
#     segmented_folder = f"{pdf_basename}_segmented"

#     # Create the folder for segmented PDFs
#     os.makedirs(segmented_folder, exist_ok=True)

#     # Open the PDF file
#     with open(pdf_path, 'rb') as pdf_file:
#         pdf_reader = PdfReader(pdf_file)
#         total_pages = len(pdf_reader.pages)

#         current_writer = PdfWriter()
#         current_file_index = 1
#         current_output_path = os.path.join(segmented_folder, f"{pdf_basename}_{current_file_index}.pdf")

#         for page_num in range(total_pages):
#             current_writer.add_page(pdf_reader.pages[page_num])

#             # Check if the number of pages in the current writer has reached the maximum
#             if (page_num + 1) % max_pages == 0 or page_num == total_pages - 1:
#                 # Save the current file
#                 with open(current_output_path, 'wb') as output_file:
#                     current_writer.write(output_file)

#                 print(f"Saved: {current_output_path}")

#                 # Start a new file
#                 current_file_index += 1
#                 current_output_path = os.path.join(segmented_folder, f"{pdf_basename}_{current_file_index}.pdf")
#                 current_writer = PdfWriter()

#     print(f"PDF segmented into {current_file_index - 1} file(s) in folder: {segmented_folder}")


def parse_date(date_str):
    try:
        return datetime.strptime(date_str, "%m/%d/%Y")
    except ValueError:
        return datetime.min  # Fallback for invalid dates


def extract_text_from_selected_pages(pdf_path, selected_pages):
    extracted_text = ""

    # Sort the selected pages to optimize page range extraction
    selected_pages = sorted(set(selected_pages))  # Ensure no duplicates and sort pages

    # Convert only the required pages to images
    for page_number in selected_pages:
        try:
            # Convert the specific page to an image (1-indexed)
            images = convert_from_path(pdf_path, first_page=page_number, last_page=page_number)

            # Process the single image returned for this page
            for page_image in images:
                print(f"Processing page {page_number} using PyTesseract")

                # Perform OCR on the image
                ocr_text = pytesseract.image_to_string(page_image)
                # extracted_text += f'Page {page_number}:\n{ocr_text}\n'
                extracted_text += f"{ocr_text}"

        except Exception as e:
            print(f"Error processing page {page_number}: {e}")

    return extracted_text


def extract_text_from_all_pages(pdf_path):
    extracted_text = ""

    try:
        # Convert all pages to images
        images = convert_from_path(pdf_path)

        # Process each page image
        for page_number, page_image in enumerate(images, start=1):
            print(f"Processing page {page_number} using PyTesseract")

            # Perform OCR on the image
            ocr_text = pytesseract.image_to_string(page_image)
            extracted_text += f"Page {page_number}:\n{ocr_text}\n"

    except Exception as e:
        print(f"Error processing PDF: {e}")

    return extracted_text


def count_lines_in_file(file_path):
    try:
        with open(file_path) as file:
            lines = file.readlines()
            return len(lines)
    except Exception as e:
        print(f"Error: {e}")
        return 0


def upload_to_gemini(path, mime_type=None):
    """Uploads the given file to Gemini.
    See https://ai.google.dev/gemini-api/docs/prompting_with_media
    """
    file = genai_client.files.upload(file=path)
    print()
    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
    return file


def wait_for_files_active(files):
    """Waits for the given files to be active.

    Some files uploaded to the Gemini API need to be processed before they can be
    used as prompt inputs. The status can be seen by querying the file's "state"
    field.

    This implementation uses a simple blocking polling loop. Production code
    should probably employ a more sophisticated approach.
    """
    print("Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai_client.files.get(name=name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(3)
            file = genai_client.files.get(name=name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
    print("...all files ready")
    print()


def segment_pdf(input_pdf, pages_per_segment=100):
    # Read the PDF file
    reader = PdfReader(input_pdf)
    total_pages = len(reader.pages)
    base_name = os.path.splitext(os.path.basename(input_pdf))[
        0
    ]  # Get the base file name without extension

    # Define the main folder path using base_name
    file_path = os.path.join(os.path.expanduser("~"), "MRRs")
    # Define the segmented folder path inside the main folder
    output_folder = os.path.join(file_path, f"{base_name}_segmented")

    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Calculate the total number of segments for formatting
    segment_count = 1
    for start_page in range(0, total_pages, pages_per_segment):
        writer = PdfWriter()
        end_page = min(
            start_page + pages_per_segment, total_pages
        )  # Ensure we don't exceed total pages

        # Add pages to the segment
        for page_num in range(start_page, end_page):
            writer.add_page(reader.pages[page_num])

        # Save the segment with leading zeros in the segment count
        output_file = os.path.join(output_folder, f"{base_name}_{segment_count:02}.pdf")
        with open(output_file, "wb") as output_pdf:
            writer.write(output_pdf)
        print(f"Saved: {output_file}")
        segment_count += 1
    print(f"Segmentation complete. Files saved in folder: {output_folder}")


def segment_pdf_locally(input_pdf, pages_per_segment=100):
    # Read the PDF file
    reader = PdfReader(input_pdf)
    total_pages = len(reader.pages)
    base_name = os.path.splitext(os.path.basename(input_pdf))[
        0
    ]  # Get the base file name without extension

    # Define the main folder path using UPLOAD_FOLDER from app config
    file_path = os.path.join(app.config["UPLOAD_FOLDER"])

    # Define the segmented folder path inside the main folder
    output_folder = os.path.join(file_path, f"{base_name}_segmented")

    # Create the output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)

    # Initialize a list to store created file paths
    created_files = []

    # Calculate the total number of segments for formatting
    segment_count = 1
    for start_page in range(0, total_pages, pages_per_segment):
        writer = PdfWriter()
        end_page = min(
            start_page + pages_per_segment, total_pages
        )  # Ensure we don't exceed total pages

        # Add pages to the segment
        for page_num in range(start_page, end_page):
            writer.add_page(reader.pages[page_num])

        # Save the segment with leading zeros in the segment count
        output_file = os.path.join(output_folder, f"{base_name}_{segment_count:02}.pdf")
        with open(output_file, "wb") as output_pdf:
            writer.write(output_pdf)

        # Add the file path to the list
        created_files.append(output_file)
        print(f"Saved: {output_file}")
        segment_count += 1

    # Sort the list of created files alphabetically
    created_files.sort()

    print(f"Segmentation complete. Files saved in folder: {output_folder}")
    print(created_files)

    # Return the sorted list of files
    return created_files


###########################
# Pages
###########################


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/upload", methods=["POST"])
def upload():
    global pdf_filepath
    global main_filename
    global num_pages

    print("Upload is pressed.")

    file = request.files["pdf"]
    if file:
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(filepath)

        # Read PDF and extract page count
        reader = PdfReader(filepath)
        num_pages = len(reader.pages)

        pdf_filepath = filepath
        main_filename = file.filename
        main_filename = str(main_filename.replace(".pdf", ""))

        print("num_pages", num_pages)
        print("filepath", filepath)
        print("main_filename", main_filename)

        print("pdf_filepathfff", pdf_filepath)

        return {"filepath": file.filename, "num_pages": num_pages}


# Function to check date format
def is_valid_date(date_str, date_format="%m/%d/%Y"):
    try:
        if date_str.strip() == "-":
            return True  # Skip validation for "-"
        datetime.strptime(date_str, date_format)
        return True
    except ValueError:
        return False


@app.route("/uploadAndCheckCSV", methods=["POST"])
def uploadAndCheckCSV():
    global txt_filepath
    global main_txt_filename

    print("Inside uploadAndCheckCSV")

    file = request.files["txt"]
    output_messages = []  # Accumulate messages here

    if file:
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(filepath)

        txt_filepath = filepath
        main_txt_filename = file.filename
        main_txt_filename = str(main_filename.replace(".txt", ""))

        # Read and validate the CSV file
        error_lines = []
        duplicate_groups = {}
        unique_rows = []

        try:
            with open(txt_filepath) as file:
                csv_reader = csv.reader(file)
                for line_number, row in enumerate(csv_reader, start=1):
                    if len(row) < 4:
                        error_lines.append((line_number, "Missing column(s)"))
                        continue

                    # Check for date validity
                    date_value = row[3]
                    if not is_valid_date(date_value):
                        error_lines.append((line_number, row))

                    # Check for duplicates (type and date)
                    duplicate_key = (
                        row[2].strip(),
                        row[3].strip(),
                    )  # Ensure both fields are stripped of whitespace
                    if duplicate_key in duplicate_groups:
                        duplicate_groups[duplicate_key].append((line_number, row))
                    else:
                        duplicate_groups[duplicate_key] = [(line_number, row)]

            # Separate duplicates and unique rows
            duplicate_rows = []
            for key, group in duplicate_groups.items():  # noqa: B007
                if len(group) > 1:
                    duplicate_rows.extend(group)
                else:
                    unique_rows.append(group[0])

            # Output the results
            if error_lines:
                output_messages.append("Lines with errors or missing columns:")
                output_messages.append("--------------------------------------")
                for error in error_lines:
                    output_messages.append(f"Line {error[0]}: {error[1]}")

            output_messages.append("")

            if duplicate_rows:
                output_messages.append("Duplicate rows detected:")
                output_messages.append("------------------------")
                grouped_duplicates = {}
                for line_number, row in duplicate_rows:
                    key = (row[2].strip(), row[3].strip())
                    if key not in grouped_duplicates:
                        grouped_duplicates[key] = []
                    grouped_duplicates[key].append((line_number, row))

                for group in grouped_duplicates.values():
                    for line_number, row in group:
                        output_messages.append(f"Line {line_number}: {row}")
                    output_messages.append("")  # Add a space between groups

            if not error_lines and not duplicate_rows:
                output_messages.append("All rows are valid and no duplicates were found.")

        except Exception as e:
            output_messages.append(f"Error reading file: {str(e)}")

    else:
        output_messages.append("No file was uploaded.")

    # Return the messages in a formatted way for a textbox
    return {"errors_and_duplicates": "\n".join(output_messages)}


# @app.route('/upload-multiple', methods=['POST'])
# def upload_multiple():
#     global pdf_filepath
#     global main_filename
#     global num_pages
#     global sorted_file_paths

#     if 'pdfs' not in request.files:
#         return jsonify({'error': 'No file part provided'}), 400

#     files = request.files.getlist('pdfs')
#     if not files:
#         return jsonify({'error': 'No files selected'}), 400

#     file_paths = []

#     for index, file in enumerate(files):
#         if file and allowed_file(file.filename):
#             filename = secure_filename(file.filename)

#             if index == 0:
#                 main_filename = filename

#             file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#             file.save(file_path)
#             file_paths.append(file_path)

#     if not file_paths:
#         return jsonify({'error': 'No valid PDF files uploaded'}), 400

#     # Sort file paths alphabetically
#     sorted_file_paths = sorted(file_paths)
#     len_files = len(sorted_file_paths)

#     # Extract file names from the paths
#     file_names_str = [file_path.split('/')[-1] for file_path in file_paths]

#     # Join file names into a single string, each on a new line
#     file_names_str_single_line = "\n".join(file_names_str)
#     print(file_names_str_single_line)

#     # Return JSON response
#     return jsonify({'paths': sorted_file_paths, 'len_files': len_files, 'lines': file_names_str_single_line}), 200


@app.route("/uploadPages", methods=["POST"])
def uploadPages():
    global txt_filepath
    global main_txt_filename

    file = request.files["txt"]
    if file:
        filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
        file.save(filepath)

        txt_filepath = filepath
        main_txt_filename = file.filename
        main_txt_filename = str(main_filename.replace(".txt", ""))

        line_count = count_lines_in_file(txt_filepath)

        return {"filepath": txt_filepath, "line_count": line_count}


@app.route("/summarize", methods=["POST"])
def summarize():
    model = request.json.get("model")
    global all_data
    global manual_intervention

    # Initialize default values
    summaryText = "No Summary Available"
    summaryDate = "No Date"
    summaryTitle = "No Title"
    big_text = "No Data Processed"  # noqa: F841
    # Initialized here too so the return below is safe even if the try block
    # raises before this variable is assigned on the success path.
    big_text_to_show_only = ""

    try:
        with open(txt_filepath) as file:
            for line in file:
                # Strip any leading/trailing whitespace and split the line into parts
                line = line.strip()
                if not line:  # Skip empty lines
                    continue

                values = line.split(",")

                # Validate that there are exactly four values
                if len(values) != 6:
                    print(f"Invalid line format: {line}")
                    continue

                # Parse the values into integers
                try:
                    start_page = int(values[0])
                    end_page = int(values[1])
                    document_type = int(values[2])
                    document_date = values[3]
                    doi_from_txt = values[4]
                    manual_intervention = values[5]
                except ValueError:
                    print(f"Error parsing numbers in line: {line}")
                    continue

                # Process the values (example: print them)
                print(
                    f"Start Page: {start_page}, End Page: {end_page}, Document Type: {document_type}, date: {document_date}"
                )

                # Add additional processing logic here as needed
                selected_pages = []
                for i in range(start_page, end_page + 1):
                    selected_pages.append(i)

                option = document_type

                if option == 1:
                    system_msg = prompts["category_01"]
                elif option == 2:
                    system_msg = prompts["category_02"]
                elif option == 3:
                    system_msg = prompts["category_03"]
                elif option == 4:
                    system_msg = prompts["category_04"]
                elif option == 5:
                    system_msg = prompts["category_05"]
                elif option == 6:
                    system_msg = prompts["category_06"]
                elif option == 7:
                    system_msg = prompts["category_07"]
                elif option == 8:
                    system_msg = prompts["category_08"]
                elif option == 9:
                    system_msg = prompts["category_09"]
                elif option == 10:
                    system_msg = prompts["category_10"]
                elif option == 11:
                    system_msg = prompts["category_11"]
                elif option == 12:
                    system_msg = prompts["category_12"]
                elif option == 13:
                    system_msg = prompts["category_13"]
                elif option == 14:
                    system_msg = prompts["category_14"]
                elif option == 100:
                    system_msg = prompts["category_100"]
                else:
                    system_msg = prompts["category_100"]

                print("pdf_filepath:", pdf_filepath)

                text_to_summarize = extract_text_from_selected_pages(pdf_filepath, selected_pages)
                print("Text to Summarize:")
                print(text_to_summarize)

                completion = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": [{"type": "text", "text": f"{system_msg}"}]},
                        {
                            "role": "user",
                            "content": [{"type": "text", "text": f"{text_to_summarize}"}],
                        },
                    ],
                    temperature=0.8,
                    max_tokens=2048,
                    top_p=1,
                    frequency_penalty=0,
                    presence_penalty=0,
                    response_format={"type": "text"},
                )

                output = completion.choices[0].message.content

                completion3 = client.chat.completions.create(
                    model=model,
                    messages=[
                        {
                            "role": "system",
                            "content": [
                                {
                                    "type": "text",
                                    "text": 'You are an intelligent assistant tasked with extracting the **title** of the document and the **entity responsible for the encounter**. Follow these instructions:\n\n1. **Title Extraction**: \n   - Accurately extract the title of the document if it is explicitly clear. \n   - If the title is not exactly stated, try to infer it from the context of the document; For example, "PT Progress Note", "Office Visit", "Hospital Discharge". The title can be at the top of the document, or towards the end of the document.\n   - If the title cannot be inferred, respond with `" unknown"`. \n   \n2. **Name of Entity Responsible for the Encounter**: \n   - Identify the specific entity responsible for the encounter, which must be the name of the person or the entity. \n   - If available, use the name found in the signature section towards the end of the document to identify the entity responsible for the encounter, or at the top of the document. \n   - Only return the name of the entity that directly conducted the encounter, even if multiple names are mentioned in the text.\n   - Do not return the name of the entity that referred to this encounter or the referral provider. \n   - If no entity name is available, return `"Unknown"`.\n\n3. **Output Format**: \n   - Return the results in a single line, separated by a dash (-):  \n     `[Title] - [Name of Responsible for Encounter]`.\n   - Do not include comma ever in the title. All separations should be done with a dash.\n\n4. **Do Not Add Commentary**: \n   - Do not include explanations, context, or additional text. Return only the extracted information in the required format. ',
                                }
                            ],
                        },
                        {
                            "role": "user",
                            "content": [{"type": "text", "text": f"{text_to_summarize}"}],
                        },
                    ],
                    temperature=0.8,
                    max_tokens=2048,
                    top_p=1,
                    frequency_penalty=0,
                    presence_penalty=0,
                    response_format={"type": "text"},
                )

                output_title = completion3.choices[0].message.content

                if doi_from_txt == "-":
                    doi_final = ""
                else:
                    doi_final = f"**DOI**:{doi_from_txt},"

                text_to_add_diag = ""
                if option == 3:
                    text_to_add_diag = " [Diagnostic Study]"
                else:
                    text_to_add_diag = ""

                text_to_add_manual_intervention = ""
                if manual_intervention == "x" or manual_intervention == "X":
                    text_to_add_manual_intervention = "[ManualCheck] "
                else:
                    text_to_add_manual_intervention = ""
                print("text_to_add_manual_intervention", text_to_add_manual_intervention)

                output_dict = {
                    "summaryDate": document_date,
                    "summaryTitle": text_to_add_manual_intervention
                    + output_title
                    + text_to_add_diag
                    + f" (Pages {start_page}-{end_page})",
                    # "summaryTitle": output_title + text_to_add_diag + f' (Pages {start_page}-{end_page})',
                    "manualCheck": text_to_add_manual_intervention,
                    "summaryText": f"{doi_final} {output}",
                }

                all_data.append(output_dict)
                with open("all_data_temp.txt", "w") as file:
                    file.write(str(all_data))
                print("all dataaaaaaaaaaaaaa", all_data)

            big_text_to_show_only = ""

            for item in all_data:
                # Extract values
                summaryDate = item.get("summaryDate", "No Date")
                manualCheck = item.get("manualCheck", "-")  # noqa: F841
                summaryTitle = item.get("summaryTitle", "No Title")
                summaryText = item.get("summaryText", "No Output")

                # big_text_to_show_only += f"_{summaryDate}_\n{summaryTitle}\n{summaryText}\n\n{manualCheck}\n\n"
                big_text_to_show_only += f"_{summaryDate}_\n{summaryTitle}\n{summaryText}\n\n"

            # print(big_text)

    except FileNotFoundError:
        print(f"File not found: {txt_filepath}")
        big_text_to_show_only = (
            f"ERROR: page-range file not found ({txt_filepath}). Upload the CSV/TXT first."
        )
    except Exception as e:
        print(f"An error occurred: {e}")
        big_text_to_show_only = f"ERROR during summarization: {e}"

    return {
        "summaryText": summaryText,
        "summaryDate": summaryDate,
        "summaryTitle": summaryTitle,
        "big_text": big_text_to_show_only,
    }  # we are only using big text to show though


@app.route("/getDiagOpRep", methods=["POST"])
def getDiagOpRep():
    model = request.json.get("model")  # noqa: F841
    global all_data
    global manual_intervention

    # File paths
    csv_file = txt_filepath
    pdf_file = pdf_filepath

    output_folder = os.path.join(
        os.path.expanduser("~"), "MRRs", main_filename + "_diag_and_op_reports"
    )
    final_pdf_path = os.path.join(output_folder, main_filename + "_diag_and_op_reports.pdf")

    # Create output directory if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # List to keep track of extracted PDFs
    extracted_pdfs = []

    # Process CSV
    with open(csv_file, newline="", encoding="utf-8") as file:
        reader = csv.reader(file)

        for row in reader:
            if len(row) >= 3 and row[2].strip() in ["3", "8"]:  # Check if 3rd value is "3" or "8"
                try:
                    start_page = int(row[0]) - 1  # Convert to 0-based index
                    end_page = int(row[1])  # End page (inclusive)

                    # Extract pages from PDF
                    with open(pdf_file, "rb") as pdf:
                        reader = PdfReader(pdf)
                        writer = PdfWriter()

                        for page_num in range(start_page, end_page):
                            writer.add_page(reader.pages[page_num])

                        # Save extracted PDF
                        output_pdf_path = os.path.join(
                            output_folder, f"segment_{start_page + 1}_to_{end_page}.pdf"
                        )
                        with open(output_pdf_path, "wb") as output_pdf:
                            writer.write(output_pdf)

                        extracted_pdfs.append(output_pdf_path)
                        print(
                            f"Extracted pages {start_page + 1} to {end_page} into {output_pdf_path}"
                        )

                except Exception as e:
                    print(f"Error processing row {row}: {e}")

    # Combine all extracted PDFs into a single PDF
    if extracted_pdfs:
        final_writer = PdfWriter()

        for pdf_path in extracted_pdfs:
            with open(pdf_path, "rb") as pdf_file:
                reader = PdfReader(pdf_file)
                for page in reader.pages:
                    final_writer.add_page(page)

        # Save combined PDF
        with open(final_pdf_path, "wb") as final_pdf:
            final_writer.write(final_pdf)

        print(f"Combined all segments into {final_pdf_path}")

        # Delete individual extracted PDFs
        for pdf_path in extracted_pdfs:
            os.remove(pdf_path)
            print(f"Deleted {pdf_path}")

    print("Process completed successfully.")
    return {"summaryText": "success"}  # we are only using big text to show though


@app.route("/getDepoRep", methods=["POST"])
def getDepoRep():
    model = request.json.get("model")  # noqa: F841
    global all_data
    global manual_intervention

    csv_file = txt_filepath
    pdf_file = pdf_filepath

    output_text = ""
    deposition_count = 0  # Counter for depositions

    output_folder = main_filename + "_depositions"
    # Define output directory in ~/MRRs/
    output_folder = os.path.join(os.path.expanduser("~"), "MRRs", main_filename + "_depositions")

    # Create output directory if it doesn't exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # List to keep track of extracted PDFs
    extracted_pdfs = []

    # Process CSV
    with open(csv_file, newline="", encoding="utf-8") as file:
        reader = csv.reader(file)

        for row in reader:
            if len(row) >= 3 and row[2].strip() in ["9"]:  # Check if 3rd value is "9"
                try:
                    start_page = int(row[0]) - 1  # Convert to 0-based index
                    end_page = int(row[1])  # End page (inclusive)

                    # Extract pages from PDF
                    with open(pdf_file, "rb") as pdf:
                        reader = PdfReader(pdf)
                        writer = PdfWriter()

                        for page_num in range(start_page, end_page):
                            writer.add_page(reader.pages[page_num])

                        # Save extracted PDF
                        output_pdf_path = os.path.join(
                            output_folder,
                            f"{main_filename}_Deposition_{start_page + 1}_to_{end_page}.pdf",
                        )
                        with open(output_pdf_path, "wb") as output_pdf:
                            writer.write(output_pdf)

                        extracted_pdfs.append(output_pdf_path)
                        deposition_count += 1  # Increment deposition count
                        print(
                            f"Extracted pages {start_page + 1} to {end_page} into {output_pdf_path}"
                        )

                except Exception as e:
                    print(f"Error processing row {row}: {e}")

    print("Process completed successfully.")
    output_text = f"Total Depositions: {deposition_count}\n"

    return {"summaryText": output_text}  # we are only using big text to show though


@app.route("/exportresultstoCSV", methods=["POST"])
def exportresultstoCSV():
    global main_filename

    try:
        # Get the JSON data from the request
        data = request.get_json()

        # Extract the text content and strip unwanted quotes
        csv_content = data.get("TXTText", "").strip('"')
        print(csv_content)

        # Prepare the file path
        file_path = os.path.join(os.path.expanduser("~"), "MRRs", main_filename + ".csv")
        print(main_filename)

        # Write the content to the CSV file without using csv.writer
        with open(file_path, "w", encoding="utf-8") as csv_file:
            csv_file.write(csv_content)

        # Return a success response
        return jsonify({"message": "Content saved to CSV file successfully!"}), 200
    except Exception as e:
        # Handle any errors
        return jsonify({"error": str(e)}), 500


# @app.route('/exportresultstoTXT', methods=['POST'])
# def exportresultstoTXT():
#     global main_filename

#     try:
#         # Get the JSON data from the request
#         data = request.get_json()

#         # Extract the text content
#         txt_content = data.get("TXTText", "")

#         file_path = os.path.join(os.path.expanduser("~"), "MRRs", main_filename + ".txt")
#         print(main_filename)

#         # Write the content to the file
#         with open(file_path, "w", encoding="utf-8") as txt_file:
#             txt_file.write(txt_content)

#         # Return a success response
#         return jsonify({"message": "Content saved to TXT file successfully!"}), 200
#     except Exception as e:
#         # Handle any errors
#         return jsonify({"error": str(e)}), 500


@app.route("/exportresultstoword", methods=["POST"])
def exportresultstoword():
    global all_data
    global main_filename
    global patientNameGlobal

    patientName = request.json.get("patientName")
    patientdob = request.json.get("patientdob")
    QMEorAME = request.json.get("QMEorAME")
    lawfirm = request.json.get("lawfirm")

    # Check if main_filename is None and set a default value if needed
    if main_filename is None:
        main_filename = "default_filename"

    # all_data = sorted(all_data, key=lambda x: datetime.strptime(x['summaryDate'], '%m/%d/%Y'))
    def safe_date_parse(entry):
        try:
            return datetime.strptime(entry.get("summaryDate", ""), "%m/%d/%Y")
        except ValueError:
            return datetime.min  # Assign a very early date as fallback

    all_data = sorted(all_data, key=safe_date_parse)

    # Create a new Word document and add the content
    doc = Document()

    # HEADER
    section = doc.sections[0]  # Access the first section of the document
    header = section.header  # Access the header of the section
    header_paragraph = header.add_paragraph(
        "RE: " + patientName + "\n" + patientdob + "\n" + "Page "
    )
    for run in header_paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)  # Optional: Set a font size for the header

    # Add empty lines for spacing
    doc.add_paragraph("")

    # TITLE
    # Add a title to the document with specific formatting
    title = doc.add_paragraph(QMEorAME)
    title_format = title.runs[0]  # Access the first run in the paragraph
    title_format.bold = True  # Make the text bold
    title_format.underline = True  # Underline the text
    title_format.font.size = Pt(12)  # Optional: Set the font size
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the title
    title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    # Add empty lines for spacing
    doc.add_paragraph("")

    # Medical Record Review
    second_title = doc.add_paragraph("Medical Record Review")
    second_title_format = second_title.runs[0]  # Access the first run in the paragraph
    second_title_format.bold = True  # Make the text bold
    second_title_format.underline = True  # Underline the text
    second_title_format.font.size = Pt(12)  # Optional: Set the font size
    second_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    second_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    intro_text = (
        "I have received "
        + str(num_pages)
        + " pages of medical records from "
        + lawfirm
        + ". I have reviewed all of the pages  received and my opinion is based upon such received records."
    )

    second_intro_text = "The following is a summary of those records:"
    this_concludes_text = "This concludes the review of submitted records."
    # text_1_after_mrr = "I have reviewed the documents on " + patientName + ", which we received on Xxxxxxx XX, 20XX. Exactly XX pages of documents are received. $$$Out of the stack, exactly XX pages are remarked upon, as XX pages are other documents such as:$$$\n"
    # text_2_after_mrr = "\t1)\t Records from State of California"
    # text_3_after_mrr = "\t2)\t Records from various sources"
    # text_other_documents     = "\t\ta)\t cover page\n\t\tb)\t e-mail\n\t\tc)\t cover letter\n\t\td)\t schedule of records\n\t\te)\t proof of service"
    # duplicate_copies         = "$$$There are also XX pages of duplicate copies of reports from XXXXX XXXXXX.$$$\n"
    # end_of_mrr_text          = "XX: XXX/XXX\n20XX-XXXXXX\nXXX"

    # First Line
    third_title = doc.add_paragraph(intro_text)
    third_title_format = third_title.runs[0]  # Access the first run in the paragraph
    third_title_format.bold = False  # Make the text bold
    third_title_format.underline = False  # Underline the text
    third_title_format.font.size = Pt(12)  # Optional: Set the font size
    third_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    third_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    # Second Line
    fourth_title = doc.add_paragraph(second_intro_text)
    fourth_title_format = fourth_title.runs[0]  # Access the first run in the paragraph
    fourth_title_format.bold = True  # Make the text bold
    fourth_title_format.underline = False  # Underline the text
    fourth_title_format.font.size = Pt(12)  # Optional: Set the font size
    fourth_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    fourth_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    big_text = ""
    for entry in all_data:
        # big_text += f"_{entry['summaryDate']}_\t****{entry['summaryTitle']}****: {entry['summaryText']}\n{entry['manualCheck']}"
        big_text += (
            f"_{entry['summaryDate']}_\t****{entry['summaryTitle']}****: {entry['summaryText']}"
        )

        # print('')
        # print('Big TEXTTT')
        # print(big_text)
        # print('')

        main_paragraph = doc.add_paragraph(big_text)
        big_text = ""

    for run in main_paragraph.runs:  # Ensure all runs in the paragraph use Times New Roman
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)  # Optional: Set a standard font size for the content

    nine_title = doc.add_paragraph(this_concludes_text)  # noqa: F841
    nine_title_format = fourth_title.runs[0]  # Access the first run in the paragraph
    nine_title_format.bold = False  # Make the text bold
    nine_title_format.underline = False  # Underline the text
    nine_title_format.font.size = Pt(12)  # Optional: Set the font size
    nine_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    nine_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    file_path_int = os.path.join(os.path.expanduser("~"), "MRRs", main_filename + "_int.docx")
    file_path_rep = os.path.join(os.path.expanduser("~"), "MRRs", main_filename + "_rep.docx")

    # file_path = os.path.join(os.path.expanduser("~"), "MRRs", main_filename + ".docx")
    # file_path = os.path.join("/home", "MRRs", main_filename + ".docx")

    # doc.save(file_path)

    # # Return the document as a downloadable file with explicit headers
    # response = send_file(
    #     file_path,
    #     as_attachment=True,
    #     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    # )

    doc.save(file_path_int)

    # Return the document as a downloadable file with explicit headers
    response_int = send_file(
        file_path_int,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    doc.save(file_path_rep)

    # Return the document as a downloadable file with explicit headers
    response_rep = send_file(
        file_path_rep,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response_int.headers["Content-Disposition"] = "attachment; filename=summaries.docx"
    response_rep.headers["Content-Disposition"] = "attachment; filename=summaries.docx"

    return response_int


@app.route("/exportResultsToWordFileIndivRecords", methods=["POST"])
def exportResultsToWordFileIndivRecords():
    global all_data
    global main_filename
    global patientNameGlobal

    patientName = request.json.get("patientName")
    patientdob = request.json.get("patientdob")
    QMEorAME = request.json.get("QMEorAME")
    lawfirm = request.json.get("lawfirm")

    # Check if main_filename is None and set a default value if needed
    if main_filename is None:
        main_filename = "default_filename"

    # all_data = sorted(all_data, key=lambda x: datetime.strptime(x['summaryDate'], '%m/%d/%Y'))

    def safe_date_parse(entry):
        try:
            return datetime.strptime(entry.get("summaryDate", ""), "%m/%d/%Y")
        except ValueError:
            return datetime.min  # Assign a very early date as fallback

    all_data = sorted(all_data, key=safe_date_parse)

    # Create a new Word document and add the content
    doc = Document()

    # HEADER
    section = doc.sections[0]  # Access the first section of the document
    header = section.header  # Access the header of the section
    header_paragraph = header.add_paragraph(
        "RE: " + patientName + "\n" + patientdob + "\n" + "Page "
    )
    for run in header_paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)  # Optional: Set a font size for the header

    # Add empty lines for spacing
    doc.add_paragraph("")

    # TITLE
    # Add a title to the document with specific formatting
    title = doc.add_paragraph(QMEorAME)
    title_format = title.runs[0]  # Access the first run in the paragraph
    title_format.bold = True  # Make the text bold
    title_format.underline = True  # Underline the text
    title_format.font.size = Pt(12)  # Optional: Set the font size
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the title
    title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    # Add empty lines for spacing
    doc.add_paragraph("")

    # Medical Record Review
    second_title = doc.add_paragraph("Medical Record Review")
    second_title_format = second_title.runs[0]  # Access the first run in the paragraph
    second_title_format.bold = True  # Make the text bold
    second_title_format.underline = True  # Underline the text
    second_title_format.font.size = Pt(12)  # Optional: Set the font size
    second_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    second_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    intro_text = (
        "I have received "
        + str(num_pages)
        + " pages of medical records from "
        + lawfirm
        + ". I have reviewed all of the pages  received and my opinion is based upon such received records."
    )

    second_intro_text = "The following is a summary of those records:"
    this_concludes_text = "This concludes the review of submitted records."
    # text_1_after_mrr = "I have reviewed the documents on " + patientName + ", which we received on Xxxxxxx XX, 20XX. Exactly XX pages of documents are received. $$$Out of the stack, exactly XX pages are remarked upon, as XX pages are other documents such as:$$$\n"
    # text_2_after_mrr = "\t1)\t Records from State of California"
    # text_3_after_mrr = "\t2)\t Records from various sources"
    # text_other_documents     = "\t\ta)\t cover page\n\t\tb)\t e-mail\n\t\tc)\t cover letter\n\t\td)\t schedule of records\n\t\te)\t proof of service"
    # duplicate_copies         = "$$$There are also XX pages of duplicate copies of reports from XXXXX XXXXXX.$$$\n"
    # end_of_mrr_text          = "XX: XXX/XXX\n20XX-XXXXXX\nXXX"

    # First Line
    third_title = doc.add_paragraph(intro_text)
    third_title_format = third_title.runs[0]  # Access the first run in the paragraph
    third_title_format.bold = False  # Make the text bold
    third_title_format.underline = False  # Underline the text
    third_title_format.font.size = Pt(12)  # Optional: Set the font size
    third_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    third_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    # Second Line
    fourth_title = doc.add_paragraph(second_intro_text)
    fourth_title_format = fourth_title.runs[0]  # Access the first run in the paragraph
    fourth_title_format.bold = True  # Make the text bold
    fourth_title_format.underline = False  # Underline the text
    fourth_title_format.font.size = Pt(12)  # Optional: Set the font size
    fourth_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    fourth_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    big_text = ""
    for entry in all_data:
        # big_text += f"_{entry['summaryDate']}_\t****{entry['summaryTitle']}****: {entry['summaryText']}\n{entry['manualCheck']}"
        big_text += (
            f"_{entry['summaryDate']}_\t****{entry['summaryTitle']}****: {entry['summaryText']}"
        )

        # print('')
        # print('Big TEXTTT')
        # print(big_text)
        # print('')

        main_paragraph = doc.add_paragraph(big_text)
        big_text = ""

    for run in main_paragraph.runs:  # Ensure all runs in the paragraph use Times New Roman
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)  # Optional: Set a standard font size for the content

    nine_title = doc.add_paragraph(this_concludes_text)  # noqa: F841
    nine_title_format = fourth_title.runs[0]  # Access the first run in the paragraph
    nine_title_format.bold = False  # Make the text bold
    nine_title_format.underline = False  # Underline the text
    nine_title_format.font.size = Pt(12)  # Optional: Set the font size
    nine_title_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align the title to the left
    nine_title_format.font.name = "Times New Roman"  # Set the font to Times New Roman

    file_path = os.path.join(os.path.expanduser("~"), "MRRs", patientName + " - MRR.docx")
    doc.save(file_path)

    # Return the document as a downloadable file with explicit headers
    response = send_file(
        file_path,
        as_attachment=True,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    response.headers["Content-Disposition"] = "attachment; filename=summaries.docx"

    return response


@app.route("/reset", methods=["POST"])
def reset():
    # Clear the session data
    # session.clear()
    global pdf_filepath
    global main_filename
    global all_data
    global pages_not_counting

    pdf_filepath = None
    main_filename = "summary"
    all_data = []
    pages_not_counting = 0

    return {"message": "Session data cleared successfully"}, 200


@app.route("/getpatientnameanddob", methods=["POST"])
def getpatientnameanddob():
    text_to_summarize = extract_text_from_selected_pages(pdf_filepath, [5, 15])

    completion3 = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": [
                    {
                        "type": "text",
                        "text": "You are an assistant that will extract the name of the patient and their DOB from the text and return it in a JSON format with name and dob as the keys. Make the DOB format mm/dd/yyyy",
                    }
                ],
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Extract the name of the patient and their date of birth (DOB) from this text: "
                        + text_to_summarize,
                    }
                ],
            },
        ],
        temperature=1,
        max_tokens=2048,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
        response_format={"type": "text"},
    )

    output = completion3.choices[0].message.content
    clean_response = output.replace("```json", "").replace("```", "").strip()
    print(clean_response)

    json_data = json.loads(clean_response)
    name = json_data.get("name")
    dob = json_data.get("dob")

    print(name, dob)
    return {"name": name, "dob": dob}


@app.route("/getlawfirm", methods=["POST"])
def getlawfirm():
    text_to_summarize = extract_text_from_selected_pages(pdf_filepath, [1, 7])

    completion3 = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": [
                    {
                        "type": "text",
                        "text": "You are an assistant that will extract the name of the lawyer or attorney sending the document, as well as the name of the law firm they represent",
                    }
                ],
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": "Extract the name of the attorney and the law firm it represents and return it in a JSON format with the key 'lawfirm' and the value being the name of the attorney, followed by 'from' the name of lawfirm. The name of the attorney and the law firm is the declaration page. (Note that this name is different than the doctor). Use this text: "
                        + text_to_summarize,
                    }
                ],
            },
        ],
        temperature=1,
        max_tokens=2048,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
        response_format={"type": "text"},
    )

    output = completion3.choices[0].message.content
    clean_response = output.replace("```json", "").replace("```", "").strip()
    print(clean_response)

    json_data = json.loads(clean_response)
    lawfirm = json_data.get("lawfirm")

    print(lawfirm, lawfirm)
    return {"lawfirm": lawfirm}


@app.route("/pages")
def pages():
    return render_template("pages.html")


@app.route("/pagesManual")
def pagesManual():
    return render_template("pagesManual.html")


@app.route("/pdfsegment")
def pdfsegment():
    return render_template("pdfsegment.html")


@app.route("/checkCSV")
def checkCSV():
    return render_template("checkCSV.html")


@app.route("/DiagAndOpReports")
def DiagAndOpReports():
    return render_template("DiagAndOpReports.html")


@app.route("/DepositionReports")
def DepositionReports():
    return render_template("DepositionReports.html")


@app.route("/IndividualMRR")
def IndividualMRR():
    return render_template("individual_mrr.html")


UPLOAD_BASE_DIR = "uploads"  # Base directory for all patient folders


@app.route("/create_patient_folder_indiv_mrr", methods=["POST"])
def create_patient_folder():
    """Creates a patient folder but does not upload files."""
    data = request.json
    folder_name = data.get("folder_name")
    patientName = data.get("patient_name")

    global patientNameGlobal
    patientNameGlobal = patientName

    if not folder_name:
        return jsonify({"error": "Invalid folder name"}), 400

    folder_path = os.path.join(UPLOAD_BASE_DIR, folder_name)

    try:
        os.makedirs(folder_path, exist_ok=True)  # Create the directory if it doesn't exist
        return jsonify(
            {"message": f"Folder '{folder_name}' created successfully", "folder_path": folder_path}
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/upload_files", methods=["POST"])
def upload_files():
    """Uploads files to the already created patient folder."""
    global indiv_mrr_folder_path
    patient_folder = request.form.get("folder_name")  # Get the folder name from form data

    if not patient_folder:
        return jsonify({"error": "Missing patient folder name"}), 400

    folder_path = os.path.join(UPLOAD_BASE_DIR, patient_folder)
    indiv_mrr_folder_path = folder_path
    print("22", indiv_mrr_folder_path)

    if not os.path.exists(folder_path):
        return jsonify({"error": "Patient folder does not exist"}), 400

    files = request.files.getlist("pdfs")  # Retrieve multiple files from the request

    saved_files = []
    for file in files:
        if file:
            filename = file.filename  # Use original filename
            file_path = os.path.join(folder_path, filename)  # Save in the patient folder
            file.save(file_path)
            saved_files.append(filename)

    return jsonify(
        {
            "message": "Files uploaded successfully",
            "saved_files": saved_files,
            "folder_path": folder_path,
        }
    )


UPLOAD_BASE_DIR = "uploads"  # Base directory for patient folders


@app.route("/compute_page_ranges", methods=["POST"])
def compute_page_ranges():
    """Compute page ranges for PDFs and merge them into a single file."""

    data = request.json
    folder_name = data.get("folder_name")

    if not folder_name:
        return jsonify({"error": "Missing folder name"}), 400

    folder_path = os.path.join(UPLOAD_BASE_DIR, folder_name)

    if not os.path.exists(folder_path):
        return jsonify({"error": f"Folder '{folder_path}' does not exist"}), 400

    pdf_files = sorted([f for f in os.listdir(folder_path) if f.endswith(".pdf")])
    if not pdf_files:
        return jsonify({"error": "No PDF files found in folder"}), 400

    page_ranges = []
    current_page = 1
    merger = PdfWriter()

    for pdf_file in pdf_files:
        pdf_path = os.path.join(folder_path, pdf_file)

        try:
            pdf_reader = PdfReader(pdf_path)
            num_pages = len(pdf_reader.pages)
        except Exception as e:
            print(f"Error reading {pdf_file}: {e}")
            num_pages = 0

        if num_pages > 0:
            start_page = current_page
            end_page = current_page + num_pages - 1
            page_ranges.append(f"{start_page}-{end_page}")
            current_page = end_page + 1  # Move to the next range

            # Add PDF to merger
            merger.append(pdf_path)

    # Save the merged PDF as 'AAA.pdf' inside the folder
    # merged_pdf_path = os.path.join(folder_path, "AAA.pdf")
    merged_pdf_path = os.path.join(
        os.path.expanduser("~"), "MRRs", patientNameGlobal + " - AGGREGATED_RECORDS.pdf"
    )
    merger.write(merged_pdf_path)
    merger.close()

    return jsonify({"status": "success", "page_ranges": page_ranges, "merged_pdf": merged_pdf_path})


@app.route("/summarize_indiv_record", methods=["POST"])
def summarize_indiv_record():
    print("we are here 1")

    global indiv_mrr_folder_path
    global all_data
    global manual_intervention
    all_data = []

    model = "gpt-4o-mini"

    """Iterates through the rows, retrieves metadata from input fields, and prints record details."""

    data = request.json  # Get JSON data from request
    print(data)
    folder_name = data.get("folder_name")  # noqa: F841
    records = data.get("records", [])

    if not records:
        return jsonify({"error": "No records received"}), 400

    summary_results = []  # noqa: F841

    for record in records:
        print("Starting RECORD")
        print("--------------------")
        filename = record.get("filename", "Unknown")
        category = record.get("category", "100")
        encounter_date = record.get("encounter_date", "01/01/1900")
        injury_date = record.get("injury_date", "01/01/1900")
        manual_review = record.get("manual_review", "-")
        pages = record.get("pages", "-")

        # Construct full file path
        full_path = os.path.join(indiv_mrr_folder_path, filename)
        print("fl", full_path)

        # Print record details for debugging
        print(f"Processing record: {full_path}")
        print(
            f"Category: {category}, Encounter Date: {encounter_date}, Injury Date: {injury_date}, Manual Review: {manual_review}, Pages: {pages}"
        )

        summaryText = "No Summary Available"
        summaryDate = "No Date"
        summaryTitle = "No Title"
        big_text = "No Data Processed"  # noqa: F841

        option = category

        try:
            if option == 1 or "1":
                system_msg = prompts["category_01"]
            elif option == 2 or "2":
                system_msg = prompts["category_02"]
            elif option == 3 or "3":
                print("here")
                system_msg = prompts["category_03"]
            elif option == 4 or "4":
                system_msg = prompts["category_04"]
            elif option == 5 or "5":
                system_msg = prompts["category_05"]
            elif option == 6 or "6":
                system_msg = prompts["category_06"]
            elif option == 7 or "7":
                system_msg = prompts["category_07"]
            elif option == 8 or "8":
                system_msg = prompts["category_08"]
            elif option == 9 or "9":
                system_msg = prompts["category_09"]
            elif option == 10 or "10":
                system_msg = prompts["category_10"]
            elif option == 11 or "11":
                system_msg = prompts["category_11"]
            elif option == 12 or "12":
                system_msg = prompts["category_12"]
            elif option == 13 or "13":
                system_msg = prompts["category_13"]
            elif option == 14 or "14":
                system_msg = prompts["category_14"]
            elif option == 100 or "100":
                system_msg = prompts["category_100"]
            else:
                system_msg = prompts["category_100"]

            print("pdf_filepath:", full_path)

            text_to_summarize = extract_text_from_all_pages(full_path)
            # print('Text to Summarize:')
            print(text_to_summarize)
        except:  # noqa: E722
            print("except")

        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": [{"type": "text", "text": f"{system_msg}"}]},
                {"role": "user", "content": [{"type": "text", "text": f"{text_to_summarize}"}]},
            ],
            temperature=0.8,
            max_tokens=2048,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            response_format={"type": "text"},
        )

        output = completion.choices[0].message.content
        # print('output_main_summary', output)

        completion3 = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": [
                        {
                            "type": "text",
                            "text": 'You are an intelligent assistant tasked with extracting the **title** of the document and the **entity responsible for the encounter**. Follow these instructions:\n\n1. **Title Extraction**: \n   - Accurately extract the title of the document if it is explicitly clear. \n   - If the title is not exactly stated, try to infer it from the context of the document; For example, "PT Progress Note", "Office Visit", "Hospital Discharge". The title can be at the top of the document, or towards the end of the document.\n   - If the title cannot be inferred, respond with `" unknown"`. \n   \n2. **Name of Entity Responsible for the Encounter**: \n   - Identify the specific entity responsible for the encounter, which must be the name of the person or the entity. \n   - If available, use the name found in the signature section towards the end of the document to identify the entity responsible for the encounter, or at the top of the document. \n   - Only return the name of the entity that directly conducted the encounter, even if multiple names are mentioned in the text.\n   - Do not return the name of the entity that referred to this encounter or the referral provider. \n   - If no entity name is available, return `"Unknown"`.\n\n3. **Output Format**: \n   - Return the results in a single line, separated by a dash (-):  \n     `[Title] - [Name of Responsible for Encounter]`.\n   - Do not include comma ever in the title. All separations should be done with a dash.\n\n4. **Do Not Add Commentary**: \n   - Do not include explanations, context, or additional text. Return only the extracted information in the required format. ',
                        }
                    ],
                },
                {"role": "user", "content": [{"type": "text", "text": f"{text_to_summarize}"}]},
            ],
            temperature=0.8,
            max_tokens=2048,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            response_format={"type": "text"},
        )

        output_title = completion3.choices[0].message.content
        # print('output_title', output_title)

        if injury_date == "-" or injury_date == "":
            doi_final = ""
        else:
            doi_final = f"**DOI**:{injury_date},"

        text_to_add_diag = ""
        if category == 3 or category == "3":
            text_to_add_diag = " [Diagnostic Study]"
        else:
            text_to_add_diag = ""

        text_to_add_manual_intervention = ""
        if manual_review == "x" or manual_review == "X":
            text_to_add_manual_intervention = "[ManualCheck] "
        else:
            text_to_add_manual_intervention = ""
        # print('text_to_add_manual_intervention', text_to_add_manual_intervention)

        # filename = record.get("filename", "Unknown")
        # category = record.get("category", "100")
        # encounter_date = record.get("encounter_date", "01/01/1900")
        # injury_date = record.get("injury_date", "01/01/1900")
        # manual_review = record.get("manual_review", "-")

        output_dict = {
            "summaryDate": encounter_date,
            # "summaryTitle": text_to_add_manual_intervention + output_title + text_to_add_diag,
            "summaryTitle": text_to_add_manual_intervention
            + output_title
            + text_to_add_diag
            + f" (Pages: {pages})",
            "manualCheck": text_to_add_manual_intervention,
            "summaryText": f"{doi_final} {output}",
        }

        all_data.append(output_dict)
        with open("all_data_temp.txt", "w") as file:
            file.write(str(all_data))
        # print('all dataaaaaaaaaaaaaa', all_data)
        print("")

        big_text_to_show_only = ""

        for item in all_data:
            # Extract values
            summaryDate = item.get("summaryDate", "No Date")
            manualCheck = item.get("manualCheck", "-")  # noqa: F841
            summaryTitle = item.get("summaryTitle", "No Title")
            summaryText = item.get("summaryText", "No Output")

            # big_text_to_show_only += f"_{summaryDate}_\n{summaryTitle}\n{summaryText}\n\n{manualCheck}\n\n"
            big_text_to_show_only += f"_{summaryDate}_\n{summaryTitle}\n{summaryText}\n\n"

    print("all dataaaaaaaaaaaaaa", all_data)

    # print(big_text_to_show_only)

    # except FileNotFoundError:
    #     print(f"File not found: {txt_filepath}")
    # except Exception as e:
    #     print(f"An error occurred: {e}")

    return "S"
    # return {"summaryText": summaryText, "summaryDate": summaryDate, "summaryTitle": summaryTitle, "big_text": big_text_to_show_only} #we are only using big text to show though

    # return jsonify({"message": "Records processed successfully", "records": summary_results})


@app.route("/segmentPDF", methods=["POST"])
def segmentPDF():
    global sorted_file_paths
    global pdf_filepath

    print("inside segmented PDF")

    segment_pdf(pdf_filepath, pages_per_segment=100)
    return {"pages": "File segmentation finalyzed. You can get the files from the MRR folder."}


SEGMENTATION_PROMPT = """Title: Extract Subdocument Metadata from a PDF

I have a large PDF document containing multiple subdocuments, each of which can vary in type (e.g., diagnostic reports, doctor's notes, legal forms, etc.).
Your task is to analyze the PDF and return a structured JSON array containing key metadata for each subdocument. Use EXACTLY these keys for every element:

1) "id" (subdocument ID): A unique identifier for each subdocument (e.g., Doc1, Doc2, etc.).
2) "s" (start page): The page number where the subdocument begins (integer).
3) "e" (end page): The page number where the subdocument ends (integer).
4) "t" (title): The title or header of the subdocument, if available. Do not invent titles; if needed, infer from the document type. DO NOT use commas; convert any comma to a dash (-). For example: WORK ACTIVITY STATUS
5) "d" (date of the document/encounter): The visit/encounter date as MM/DD/YYYY, else "-". If there are several dates, pick the one labeled visit or encounter. The date can be near the signature at the end.
6) "i" (date of injury): The injury date as MM/DD/YYYY, else "-".
7) "m" (manual check): Return "x" if the document (1) has handwriting other than a signature, (2) has many checkboxes with x/ticks, (3) is a work status report, or (4) is a QME/AME report; otherwise "-".

## Guidelines for Extraction:
- Cover every page; do not skip any page.
- Use contextual clues such as headers, bold titles, or consistent formatting to identify boundaries and titles.
- Link pages together using page counts to figure out where documents start and end.
- Distinguish the document/encounter date from the injury date.
- A title can sometimes be next to a word such as 'Notes'.
- If a field is unavailable, use "-" (never None or null).
- Ignore fax/resend dates; use the encounter/visit date or the day the document was created.
- If a title contains "X vs Y", it is most likely a deposition; set "t" to "Deposition".
- If the only handwriting is a signature, return "-" not "x".
- If a page is empty, set "t" to "Empty Page".
- Do not split a single document into two, and do not merge two documents into one.
- QME/PQME/AME evaluations can be long and often quote other records; treat the entire evaluation as ONE record with the correct start and end pages.
- Different QME/PQME/AME supplemental reports are separate documents.
- Treat the first page of a document as part of that document.

Example JSON output for a 10-page PDF:
[
  {"id": "Doc1", "s": 1, "e": 5, "t": "WORK ACTIVITY STATUS", "d": "12/03/2021", "i": "05/07/2018", "m": "x"},
  {"id": "Doc2", "s": 6, "e": 10, "t": "ACUPUNCTURE THERAPY NOTES", "d": "11/11/2022", "i": "-", "m": "-"}
]

## IMPORTANT: Return ONLY the JSON array, with no markdown fences and no other explanation."""


def parse_segment_item(item):
    """Tolerantly extract one subdocument record from a Gemini JSON element.

    Handles the t/title key alias, missing keys, and type coercion so a single
    malformed element raises (to be skipped by the caller) rather than the old
    behavior of a KeyError aborting the entire batch.
    """
    title = item.get("t") or item.get("title") or "-"
    if not isinstance(title, str):
        title = str(title)
    return (
        int(item["s"]),
        int(item["e"]),
        title.strip(),
        str(item.get("d", "-")).strip(),
        str(item.get("i", "-")).strip(),
        str(item.get("m", "-")).strip(),
    )


# Automatic segmentation upload
@app.route("/getPages", methods=["POST"])
def getPages():
    global sorted_file_paths
    global pdf_filepath
    global main_filename
    global num_pages

    # Get the JSON data from the request
    data = request.json

    # Extract the pageDelimiter value, default to 100 if not provided
    page_delimiter = data.get("pageDelimiter", 100)

    # Ensure pageDelimiter is an integer
    try:
        page_delimiter = int(page_delimiter)
    except ValueError:
        page_delimiter = 100

    # Process page_delimiter as needed
    print(f"Page delimiter received: {page_delimiter}")

    pdf_size = get_pdf_size(pdf_filepath)
    pdf_pages = get_pdf_page_count(pdf_filepath)
    print(f"PDF Size: {pdf_size} MB, PDF Pages: {pdf_pages}")

    segmentation_model = "gemini-flash-latest"
    generation_config = types.GenerateContentConfig(
        temperature=0.0,
        top_p=0.95,
        top_k=40,
        response_mime_type="application/json",
        system_instruction="You are an assistant that segments a large document into subdocuments and provide their metadata.",
    )

    if pdf_size > 45 or pdf_pages > 100:
        print("PDF is large. Will segment batches.")
        sorted_file_paths = segment_pdf_locally(pdf_filepath, pages_per_segment=page_delimiter)

        offsets = {}
        current_offset = 0
        lines = []

        for pdf_path in sorted_file_paths:
            print(f"Working on: {pdf_path}")
            files = [upload_to_gemini(pdf_path, mime_type="application/pdf")]
            wait_for_files_active(files)
            print("The file has been uploaded to sucessfully.")

            print("Preparing for the AI")

            # gemini change here
            prompt = SEGMENTATION_PROMPT

            try:
                response = genai_client.models.generate_content(
                    model=segmentation_model,
                    contents=[files[0], prompt],
                    config=generation_config,
                )
                print(f"Response received: {response}")
            except Exception as e:
                print(f"An error occurred while sending the message: {e}")
                # Optional: Handle the error (e.g., return a default response or log more details)
                return {"pages": str(e)}

            print(response)
            print()
            print(response.text)
            print("Response received from server.")

            clean_response = response.text.replace("```json", "").replace("```", "").strip()
            clean_response_json = json.loads(clean_response)

            print("Response converted to JSON")

            print(clean_response_json)
            print("Response formatted.")

            for item in clean_response_json:
                try:
                    start_page, end_page, title, date, injury_date, manual_check = (
                        parse_segment_item(item)
                    )
                except (KeyError, TypeError, ValueError) as exc:
                    print(f"Skipping malformed segment: {exc}")
                    continue

                title_group = categorize_documents(title, groups)

                # Create the line
                # line = f"{start_page+current_offset},{end_page+current_offset},{title},{title_group},{date},{injury_date},{manual_check}"
                line = f"{start_page + current_offset},{end_page + current_offset},{title_group},{date},{injury_date},{manual_check}"
                # if(title != 'Empty Page'):
                lines.append(line)

            print("All lines are created for this document")
            current_offset += page_delimiter
    else:
        print("PDF is not large. Proceeding.")
        offsets = {}  # noqa: F841
        current_offset = 0
        lines = []

        print(f"Working on: {pdf_filepath}")
        files = [upload_to_gemini(pdf_filepath, mime_type="application/pdf")]
        wait_for_files_active(files)
        print("The file has been uploaded to sucessfully.")

        print("AI Process Starting ....")

        # gemini change here
        prompt = SEGMENTATION_PROMPT

        try:
            response = genai_client.models.generate_content(
                model=segmentation_model,
                contents=[files[0], prompt],
                config=generation_config,
            )
            print(f"Response received: {response}")
        except Exception as e:
            print(f"An error occurred while sending the message: {e}")
            # Optional: Handle the error (e.g., return a default response or log more details)
            return {"pages": str(e)}

        print(response)
        print()
        print(response.text)
        print("Response is received from server.")

        clean_response = response.text.replace("```json", "").replace("```", "").strip()
        clean_response_json = json.loads(clean_response)

        print("Response is converted to JSON")

        print(clean_response_json)
        print("Response is formatted.")

        for item in clean_response_json:
            try:
                start_page, end_page, title, date, injury_date, manual_check = parse_segment_item(
                    item
                )
            except (KeyError, TypeError, ValueError) as exc:
                print(f"Skipping malformed segment: {exc}")
                continue

            title_group = categorize_documents(title, groups)

            # Create the line
            # line = f"{start_page+current_offset},{end_page+current_offset},{title},{title_group},{date},{injury_date},{manual_check}"
            line = f"{start_page + current_offset},{end_page + current_offset},{title_group},{date},{injury_date},{manual_check}"
            # if(title != 'Empty Page'):
            lines.append(line)

    # Join all lines into a single result string
    result = "\n".join(lines)
    return {"pages": result}


# NOTE: manual segmentation upload
# @app.route('/getPages', methods=['POST'])
# def getPages():
#     global sorted_file_paths

#     print('File(s) uploaded!')
#     print(sorted_file_paths)

#     offsets = {}
#     current_offset = 0
#     lines = []

#     for i, pdf_path in enumerate(sorted_file_paths):
#         try:
#             reader = PdfReader(pdf_path)
#             num_pages = len(reader.pages)
#             pdf_size = get_pdf_size(pdf_path)
#             if pdf_size > 35:
#                 print('Size is large disecting it')
#         except Exception as e:
#             print(f"Error processing {pdf_path}: {e}")


#     # Calculate offsets for each document
#     for i, pdf_path in enumerate(sorted_file_paths):
#         try:
#             reader = PdfReader(pdf_path)
#             num_pages = len(reader.pages)
#             pdf_size = get_pdf_size(pdf_path)
#             if pdf_size > 35:
#                 print('Size is large disecting it')
#         except Exception as e:
#             print(f"Error processing {pdf_path}: {e}")

#         # Process each document and generate response
#         generation_config = {
#             "temperature": 1,
#             "top_p": 0.95,
#             "top_k": 40,
#             # "max_output_tokens": 999999,
#             "response_mime_type": "text/plain",
#         }

#         model = genai.GenerativeModel(
#             model_name="gemini-1.5-flash",
#             generation_config=generation_config,
#             system_instruction="You are an assistant that segments a large document into subdocuments and provide their metadata.",
#         )

#         files = [upload_to_gemini(pdf_path, mime_type="application/pdf")]
#         wait_for_files_active(files)

#         print('The file has been uploaded to sucessfully.')

#         chat_session = model.start_chat(
#             history=[
#                 {
#                     "role": "user",
#                     "parts": [files[0]],
#                 },
#             ]
#         )

#         print('Preparing for the AI')

#         prompt="Title: Extract Subdocument Metadata from a PDF\n\nI have a large PDF document containing multiple subdocuments, each of which can vary in type (e.g., diagnostic reports, doctor's notes, legal forms, etc.). Your task is to analyze the PDF and return a structured JSON output containing key metadata for each subdocument. Each subdocument may have unique formatting or structure, but your output should consistently provide the following information:\n\n1) id (subdocument ID): A unique identifier for each subdocument (e.g., Doc1, Doc2, etc.).\n2) s (start page): The page number where the subdocument begins.\n3) e (end page): The page number where the subdocument ends.\n4) t (title): The title or header of the subdocument, if available. Be very careful and do not create titles from your own end. If needed, comprehend from the document type.\n5) d (Date of the Document/Encounter): The date of the visit if applicable, else, return '-'. Convert the data to the following format: MM/DD/YYYY. In case there are several dates, pick the one with the label visit or encounter next to it. Sometimes, the date can be at the end of the document, near the signature. \n6) i (date of injury): The date of the injury mentioned in the subdocument if applicable, else, return '-' . Convert the data to the following format: MM/DD/YYYY\n7) m (To manual check or no): If the document contains (1) handwriting (other than signature), or (2) there are many boxes that contain x or ticks, or (3) is a work status report, or (4) is a QME/AME report, this should return 'x', otherwise, return '-'.\n\n## VERY IMPORTANT\n** Your output should cover all the pages. my life depends on it.**\n\n## Guidelines for Extraction:\n- Ensure the JSON output is properly structured and contains all the required fields for each subdocument.\n- Use contextual clues such as headers, bold titles, or consistent formatting patterns to identify the boundaries and titles of subdocuments.\n- Extract dates accurately, distinguishing between the date of the document/encounter and the injury date.\n- If any of the above fields are not available in a subdocument, indicate their absence with a '-' value in the JSON output. Do not write None or Null, instead write '-'\n- As mentioned, if the document contains too much handwriting and it is hard to get the text, return 'x' for the value of manual_check\n- A document might have several dates on it, sometimes from being faxed or sent again. We do not want those dates. We want the date of the encounter or the actual visit, or the day the document was originally created. This is usually the date that is found next to the wording 'visit' from the left, right, up or down to it.\n- If the title of a document contains the term Someone vs Someone than it is most probably a deposition. Deposition can be of many types, such as Zoom Deposition, or live deposition. In all cases, make the title be 'Deposition'\n- If the only handwriting in the document is a signature, then return '-' and not 'x'\n- DO NOT SKIP ANY PAGES. GO THROUGH ALL THE PAGES IN YOUR OUTPUT. MY LIFE DEPENDS ON IT.\n- If a page is empty, make the title 'Empty Page'\n- Sometimes you are segmenting a single document and returning them as two different ones. Be careful for those. Do not segment same document.\n- Be careful for medical-legal evaluations, QME, PQME, AME reports, as these records can be long and often contain summary of other records. Treat the entirety as a single record and provide the correct start and end pages.\n- Also note that different QME, PQME or AME supplemental reports are different documents and treat it as a separated segmentation.\n- Treat the first page of a document as part of the document.\n- I noticed that sometimes when the document is large, you are mistaking the pages and not including the first page of the document as part of the segmentation and are including the first page of the next document as the last page. Be careful for that.\n\nExample JSON Output:\n[\n    {\n        \"i\": \"Doc1\",\n        \"s\": 1,\n        \"e\": 5,\n        \"title\": \"Primary Care Report\",\n        \"d\": \"12/03/2021\",\n        \"i\": \"05/07/2018\",\n\t\"m\": 'x'\n    },\n    {\n        \"id\": \"Doc2\",\n        \"s\": 6,\n        \"e\": 10,\n        \"t\": \"X-Ray Report\",\n        \"d\": \"11/11/2022\",\n        \"i\": \"-\",\n        \"m\": '-'\n    }\n]\n\n## Notes:\n- If subdocuments are not clearly titled, infer the title from the first line or prominent text in the subdocument.\n- All dates should follow the format: MM/DD/YYYY.\n- Handle subdocuments of varying lengths, even if a subdocument spans only a single page.\n- Please parse the PDF carefully and ensure high accuracy in the identification of subdocuments and metadata extraction.\n- Cover all the pages. Do not skip any page.\n\n## IMPORTANT: Return only the JSON without any other explanation.",
#         try:
#             response = chat_session.send_message(prompt)
#             print(f"Response received: {response}")
#         except Exception as e:
#             print(f"An error occurred while sending the message: {e}")
#             # Optional: Handle the error (e.g., return a default response or log more details)
#             return {"pages": str(e)}

#         # print(response)
#         # print()
#         # print(response.text)
#         print('Response received from server.')

#         clean_response = response.text.replace('```json', '').replace('```', '').strip()
#         clean_response_json = json.loads(clean_response)

#         print('Response converted to JSON')

#         print(clean_response_json)
#         print('Response formatted.')

#         for item in clean_response_json:
#             subdoc_id = item["id"]
#             start_page = item["s"]
#             end_page = item["e"]
#             title = item["t"]
#             date = item["d"]
#             injury_date = item["i"]
#             manual_check = item["m"]

#             # Categorize documents
#             title_group = categorize_documents(title, groups)

#             # Create the line
#             line = f"{start_page+current_offset},{end_page+current_offset},{title},{title_group},{date},{injury_date},{manual_check}"
#             lines.append(line)

#         print('All lines are created for this document')
#         current_offset+=num_pages

#     # Join all lines into a single result string
#     result = "\n".join(lines)
#     return {"pages": result}


if __name__ == "__main__":
    # app.run(debug=True)
    app.run(host="0.0.0.0", port=5010, debug=True)
